import asyncio
import threading
import tkinter as tk
from tkinter import ttk, scrolledtext
from PIL import Image, ImageTk  # 新增，用于图像处理
import json
from crawl4ai import AsyncWebCrawler
from crawl4ai.extraction_strategy import JsonCssExtractionStrategy, LLMExtractionStrategy
from pathlib import Path
from datetime import datetime, timedelta
import base64
from urllib.parse import urlparse, urljoin
import re
import logging
import os
import sys
import webbrowser
from pydantic import BaseModel
from typing import List
import aiofiles
import aiohttp
import mimetypes
from tqdm import tqdm
import shutil
import tkinter.filedialog
import tkinter.messagebox

# Set up logging
logging.basicConfig(level=logging.INFO)



class ScrollableFrame(ttk.Frame):
    """创建一个可滚动的框架，用于放置左侧控制面板"""

    def __init__(self, container, *args, **kwargs):
        super().__init__(container, *args, **kwargs)
        
        # 创建画布和滚动条
        self.canvas = tk.Canvas(self)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = ttk.Frame(self.canvas)

        # 配置画布
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        # 绑定事件
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        # 创建窗口
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        
        # 绑定画布大小调整
        self.canvas.bind('<Configure>', self.on_canvas_configure)
        
        # 绑定鼠标滚轮事件
        self.bind_mouse_wheel(self.canvas)
        self.bind_mouse_wheel(self.scrollable_frame)
        
        # 布局
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

    def on_canvas_configure(self, event):
        """当画布大小改变时，调整内部框架的宽度"""
        self.canvas.itemconfig(self.canvas_window, width=event.width)

    def bind_mouse_wheel(self, widget):
        """绑定鼠标滚轮事件到指定部件"""
        widget.bind("<MouseWheel>", self._on_mousewheel)  # Windows
        widget.bind("<Button-4>", self._on_mousewheel)    # Linux
        widget.bind("<Button-5>", self._on_mousewheel)    # Linux
        
        # 绑定所有子部件
        for child in widget.winfo_children():
            self.bind_mouse_wheel(child)

    def _on_mousewheel(self, event):
        """处理鼠标滚轮事件"""
        if event.num == 5 or event.delta < 0:  # 向下滚动
            self.canvas.yview_scroll(1, "units")
        elif event.num == 4 or event.delta > 0:  # 向上滚动
            self.canvas.yview_scroll(-1, "units")


class CrawlerGUI:
    def __init__(self, root):
        # 配置日志
        self.setup_logging()
        
        self.root = root
        self.root.title("Tai-网页爬虫")
        self.root.geometry("1200x800")

        # 修改数据存储路径
        self.base_dir = Path("out")
        self.data_dir = self.base_dir / "data"
        self.urls_file = self.data_dir / "url_history.json"
        self.ensure_directories()

        # 存储当前运行的已保存文件路径
        self.saved_files = []

        # 创建主框架，使用网格布局
        main_frame = ttk.Frame(root, padding="5")
        main_frame.grid(row=0, column=0, sticky=(tk.N, tk.S, tk.E, tk.W))

        # 配置网格权重
        root.columnconfigure(0, weight=1)
        root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(1, weight=1)  # 修改为1，为顶部区域留出空间

        # === 创建顶部URL和爬取区域 ===
        top_frame = ttk.Frame(main_frame)
        top_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.E, tk.W), padx=5, pady=5)
        
        # URL输入区域
        url_frame = ttk.LabelFrame(top_frame, text="URL设置", padding="5")
        url_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)

        self.url_var = tk.StringVar()
        self.url_combobox = ttk.Combobox(url_frame, textvariable=self.url_var, width=80)  # 增加宽度
        self.url_combobox.pack(side=tk.LEFT, padx=(0, 5), fill=tk.X, expand=True)

        delete_btn = ttk.Button(url_frame, text="删除", command=self.delete_current_url)
        delete_btn.pack(side=tk.RIGHT)

        # 爬取按钮
        self.crawl_button = ttk.Button(top_frame, text="开始爬取", command=self.start_crawl)
        self.crawl_button.pack(side=tk.RIGHT, padx=5)

        # === 创建左侧控制面板 ===
        scrollable_frame = ScrollableFrame(main_frame)
        scrollable_frame.grid(row=1, column=0, sticky=(tk.N, tk.S, tk.W), padx=5, pady=5)

        control_frame = scrollable_frame.scrollable_frame

        # === 浏览器配置区域 ===
        browser_frame = ttk.LabelFrame(control_frame, text="浏览器配置", padding="5")
        browser_frame.pack(fill=tk.X, pady=5)

        # 浏览器类型选择
        browser_type_frame = ttk.Frame(browser_frame)
        browser_type_frame.pack(fill=tk.X)
        ttk.Label(browser_type_frame, text="浏览器类型:").pack(side=tk.LEFT)
        self.browser_type = tk.StringVar(value="chromium")
        for text, value in [("Chromium", "chromium"), ("Firefox", "firefox"), ("WebKit", "webkit")]:
            ttk.Radiobutton(browser_type_frame, text=text, variable=self.browser_type,
                            value=value).pack(side=tk.LEFT, padx=5)

        # 基本选项
        options_frame = ttk.Frame(browser_frame)
        options_frame.pack(fill=tk.X, pady=5)
        self.headless_var = tk.BooleanVar(value=True)
        self.verbose_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="无头模式",
                        variable=self.headless_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(options_frame, text="详细日志",
                        variable=self.verbose_var).pack(side=tk.LEFT, padx=5)

        # === 爬取选项区域 ===
        crawl_frame = ttk.LabelFrame(control_frame, text="爬取选项", padding="5")
        crawl_frame.pack(fill=tk.X, pady=5)

        # 字数阈值设置
        threshold_frame = ttk.Frame(crawl_frame)
        threshold_frame.pack(fill=tk.X, pady=2)
        ttk.Label(threshold_frame, text="最小字数阈值:").pack(side=tk.LEFT)
        self.word_count_var = tk.IntVar(value=10)
        ttk.Entry(threshold_frame, textvariable=self.word_count_var,
                  width=10).pack(side=tk.LEFT, padx=5)

        # 过滤选项
        filter_frame = ttk.Frame(crawl_frame)
        filter_frame.pack(fill=tk.X, pady=2)
        self.exclude_external_links_var = tk.BooleanVar(value=True)
        self.exclude_external_images_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(filter_frame, text="排除外部链接",
                        variable=self.exclude_external_links_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(filter_frame, text="排除外部图片",
                        variable=self.exclude_external_images_var).pack(side=tk.LEFT, padx=5)

        # === 输出格式区域 ===
        format_frame = ttk.LabelFrame(control_frame, text="输出格式", padding="5")
        format_frame.pack(fill=tk.X, pady=5)

        # 基本格式选择
        basic_formats_frame = ttk.LabelFrame(format_frame, text="基本格式", padding="5")
        basic_formats_frame.pack(fill=tk.X, pady=2)

        self.output_format = tk.StringVar(value="markdown")
        basic_formats = [
            ("原始HTML", "html", "保留原始HTML标记"),
            ("清理后HTML", "cleaned_html", "移除无用标签的HTML"),
            ("Markdown", "markdown", "转换为Markdown格式"),
            ("精简Markdown", "fit_markdown", "精简的Markdown格式"),
            ("纯文本", "text", "仅保留文本内容"),
            ("结构化文本", "structured_text", "保持文本结构的纯文本")
        ]

        for text, value, tooltip in basic_formats:
            frame = ttk.Frame(basic_formats_frame)
            frame.pack(fill=tk.X)
            radio = ttk.Radiobutton(frame, text=text, variable=self.output_format, value=value)
            radio.pack(side=tk.LEFT)
            ttk.Label(frame, text=f"({tooltip})", foreground="gray").pack(side=tk.LEFT, padx=5)

        # 高级格式选项
        advanced_formats_frame = ttk.LabelFrame(format_frame, text="高级格式选项", padding="5")
        advanced_formats_frame.pack(fill=tk.X, pady=2)

        # 格式化选项
        self.format_options = {
            'preserve_images': tk.BooleanVar(value=True),
            'preserve_links': tk.BooleanVar(value=True),
            'preserve_tables': tk.BooleanVar(value=True),
            'preserve_lists': tk.BooleanVar(value=True),
            'preserve_code': tk.BooleanVar(value=True),
            'preserve_headings': tk.BooleanVar(value=True),
            'preserve_emphasis': tk.BooleanVar(value=True),
            'preserve_quotes': tk.BooleanVar(value=True)
        }

        format_options_frame = ttk.Frame(advanced_formats_frame)
        format_options_frame.pack(fill=tk.X)

        # 创建两列布局
        left_options = ttk.Frame(format_options_frame)
        left_options.pack(side=tk.LEFT, fill=tk.X, expand=True)
        right_options = ttk.Frame(format_options_frame)
        right_options.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # 左侧选项
        ttk.Checkbutton(left_options, text="保留图片", 
                        variable=self.format_options['preserve_images']).pack(anchor=tk.W)
        ttk.Checkbutton(left_options, text="保留链接", 
                        variable=self.format_options['preserve_links']).pack(anchor=tk.W)
        ttk.Checkbutton(left_options, text="保留表格", 
                        variable=self.format_options['preserve_tables']).pack(anchor=tk.W)
        ttk.Checkbutton(left_options, text="保留列表", 
                        variable=self.format_options['preserve_lists']).pack(anchor=tk.W)

        # 右侧选项
        ttk.Checkbutton(right_options, text="保留代码块", 
                        variable=self.format_options['preserve_code']).pack(anchor=tk.W)
        ttk.Checkbutton(right_options, text="保留标题", 
                        variable=self.format_options['preserve_headings']).pack(anchor=tk.W)
        ttk.Checkbutton(right_options, text="保留强调", 
                        variable=self.format_options['preserve_emphasis']).pack(anchor=tk.W)
        ttk.Checkbutton(right_options, text="保留引用", 
                        variable=self.format_options['preserve_quotes']).pack(anchor=tk.W)

        # 自定义格式选项
        custom_format_frame = ttk.LabelFrame(format_frame, text="自定义格式", padding="5")
        custom_format_frame.pack(fill=tk.X, pady=2)

        self.enable_custom_format = tk.BooleanVar(value=False)
        ttk.Checkbutton(custom_format_frame, text="启用自定义格式",
                        variable=self.enable_custom_format,
                        command=self.toggle_custom_format).pack(anchor=tk.W)

        self.custom_format_text = scrolledtext.ScrolledText(custom_format_frame, height=3, state=tk.DISABLED)
        self.custom_format_text.pack(fill=tk.X, pady=2)

        # === 高级选项区域 ===
        advanced_frame = ttk.LabelFrame(control_frame, text="高级选项", padding="5")
        advanced_frame.pack(fill=tk.X, pady=5)

        # 截图和超时设置
        screenshot_frame = ttk.Frame(advanced_frame)
        screenshot_frame.pack(fill=tk.X, pady=2)
        self.screenshot_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(screenshot_frame, text="保存截图",
                        variable=self.screenshot_var).pack(side=tk.LEFT, padx=5)

        timeout_frame = ttk.Frame(advanced_frame)
        timeout_frame.pack(fill=tk.X, pady=2)
        ttk.Label(timeout_frame, text="页面加载超时 (秒):").pack(side=tk.LEFT)
        self.timeout_var = tk.IntVar(value=60)
        ttk.Entry(timeout_frame, textvariable=self.timeout_var,
                  width=10).pack(side=tk.LEFT, padx=5)

        # 反检测选项
        detection_frame = ttk.Frame(advanced_frame)
        detection_frame.pack(fill=tk.X, pady=2)
        self.simulate_user_var = tk.BooleanVar(value=True)
        self.magic_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(detection_frame, text="模拟用户行为",
                        variable=self.simulate_user_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(detection_frame, text="启用反检测",
                        variable=self.magic_var).pack(side=tk.LEFT, padx=5)

        # === 页面交互区域 ===
        interaction_frame = ttk.LabelFrame(control_frame, text="页面交互", padding="5")
        interaction_frame.pack(fill=tk.X, pady=5)

        # JavaScript代码输入
        self.enable_js_code_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(interaction_frame, text="启用 JavaScript 代码",
                        variable=self.enable_js_code_var).pack(anchor=tk.W)
        self.js_code_text = scrolledtext.ScrolledText(
            interaction_frame, height=5, state=tk.DISABLED)
        self.js_code_text.pack(fill=tk.X, pady=2)

        self.enable_js_code_var.trace('w', self.toggle_js_code)

        # 等待条件输入
        self.enable_wait_for_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(interaction_frame, text="启用等待条件",
                        variable=self.enable_wait_for_var).pack(anchor=tk.W)
        self.wait_for_text = scrolledtext.ScrolledText(
            interaction_frame, height=3, state=tk.DISABLED)
        self.wait_for_text.pack(fill=tk.X, pady=2)

        self.enable_wait_for_var.trace('w', self.toggle_wait_for)

        # === 会话管理区域 ===
        session_frame = ttk.LabelFrame(control_frame, text="会话管理", padding="5")
        session_frame.pack(fill=tk.X, pady=5)

        self.enable_session_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(session_frame, text="启用会话管理",
                        variable=self.enable_session_var).pack(anchor=tk.W)

        session_id_frame = ttk.Frame(session_frame)
        session_id_frame.pack(fill=tk.X, pady=2)
        ttk.Label(session_id_frame, text="会话 ID:").pack(side=tk.LEFT)
        self.session_id_var = tk.StringVar()
        self.session_id_entry = ttk.Entry(
            session_id_frame, textvariable=self.session_id_var, state=tk.DISABLED)
        self.session_id_entry.pack(side=tk.LEFT, padx=5)

        self.enable_session_var.trace('w', self.toggle_session_id)

        # === �����取策略区域 ===
        extraction_frame = ttk.LabelFrame(control_frame, text="提取策略", padding="5")
        extraction_frame.pack(fill=tk.X, pady=5)

        ttk.Label(extraction_frame, text="选择提取策略:").pack(anchor=tk.W)
        self.extraction_strategy_var = tk.StringVar(value="none")
        strategies = [("无", "none"), ("JsonCssExtraction", "jsoncss"), ("LLMExtraction", "llm")]
        for text, value in strategies:
            ttk.Radiobutton(extraction_frame, text=text, variable=self.extraction_strategy_var,
                            value=value, command=self.toggle_extraction_options).pack(anchor=tk.W)

        # JsonCssExtractionStrategy 配置
        self.jsoncss_config_frame = ttk.Frame(extraction_frame)
        self.jsoncss_config_frame.pack(fill=tk.X, pady=2)
        self.jsoncss_config_frame.pack_forget()

        ttk.Label(self.jsoncss_config_frame, text="Json Schema:").pack(anchor=tk.W)
        self.jsoncss_schema_text = scrolledtext.ScrolledText(
            self.jsoncss_config_frame, height=10)
        self.jsoncss_schema_text.pack(fill=tk.X, pady=2)

        # LLMExtractionStrategy 配置
        self.llm_config_frame = ttk.Frame(extraction_frame)
        self.llm_config_frame.pack(fill=tk.X, pady=2)
        self.llm_config_frame.pack_forget()

        ttk.Label(self.llm_config_frame, text="LLM 提供者:").pack(anchor=tk.W)
        self.llm_provider_var = tk.StringVar(value="ollama/nemotron")
        ttk.Entry(self.llm_config_frame, textvariable=self.llm_provider_var).pack(fill=tk.X, pady=2)

        ttk.Label(self.llm_config_frame, text="Pydantic Schema:").pack(anchor=tk.W)
        self.llm_schema_text = scrolledtext.ScrolledText(
            self.llm_config_frame, height=10)
        self.llm_schema_text.pack(fill=tk.X, pady=2)

        ttk.Label(self.llm_config_frame, text="指令 (Instruction):").pack(anchor=tk.W)
        self.llm_instruction_text = scrolledtext.ScrolledText(
            self.llm_config_frame, height=5)
        self.llm_instruction_text.pack(fill=tk.X, pady=2)

        # === 高级参数区域 ===
        advanced_params_frame = ttk.LabelFrame(control_frame, text="高���参数", padding="5")
        advanced_params_frame.pack(fill=tk.X, pady=5)

        delay_frame = ttk.Frame(advanced_params_frame)
        delay_frame.pack(fill=tk.X, pady=2)
        ttk.Label(delay_frame, text="延迟返回 HTML (秒):").pack(side=tk.LEFT)
        self.delay_var = tk.DoubleVar(value=0.0)
        ttk.Entry(delay_frame, textvariable=self.delay_var,
                  width=10).pack(side=tk.LEFT, padx=5)

        self.js_only_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(advanced_params_frame, text="仅执行 JavaScript",
                        variable=self.js_only_var).pack(anchor=tk.W, padx=5)

        # === 创建右侧结果显示区域 ===
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.grid(row=1, column=1, sticky=(tk.N, tk.S, tk.E, tk.W), padx=5, pady=5)

        # 修改文本区域的创建方式，添加滚动支持
        def create_scrolled_text(parent):
            frame = ttk.Frame(parent)
            text = scrolledtext.ScrolledText(frame, wrap=tk.WORD)
            text.pack(fill=tk.BOTH, expand=True)
            
            # 绑定鼠标滚轮事件
            text.bind("<MouseWheel>", lambda e: self._on_mousewheel(e, text))  # Windows
            text.bind("<Button-4>", lambda e: self._on_mousewheel(e, text))    # Linux
            text.bind("<Button-5>", lambda e: self._on_mousewheel(e, text))    # Linux
            
            return frame, text

        # 内容标签页
        self.content_frame, self.content_text = create_scrolled_text(self.notebook)
        self.notebook.add(self.content_frame, text="内容")

        # 媒体标签页
        self.media_frame, self.media_text = create_scrolled_text(self.notebook)
        self.notebook.add(self.media_frame, text="媒体")

        # 链接标签页
        self.links_frame, self.links_text = create_scrolled_text(self.notebook)
        self.notebook.add(self.links_frame, text="链接")

        # 文件查看标签页
        self.files_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.files_frame, text="文件查看")

        # 文件列表和图像显示框架
        files_display_frame = ttk.Frame(self.files_frame)
        files_display_frame.pack(fill=tk.BOTH, expand=True)

        # 创建带滚动条的文件列表框
        files_list_frame = ttk.Frame(files_display_frame)
        files_list_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.files_listbox = tk.Listbox(files_list_frame)
        files_scrollbar = ttk.Scrollbar(files_list_frame, orient=tk.VERTICAL, command=self.files_listbox.yview)
        self.files_listbox.configure(yscrollcommand=files_scrollbar.set)
        
        self.files_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        files_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 绑定文件列表框的鼠标滚轮事件
        self.files_listbox.bind("<MouseWheel>", lambda e: self._on_mousewheel(e, self.files_listbox))
        self.files_listbox.bind("<Button-4>", lambda e: self._on_mousewheel(e, self.files_listbox))
        self.files_listbox.bind("<Button-5>", lambda e: self._on_mousewheel(e, self.files_listbox))
        
        self.files_listbox.bind('<<ListboxSelect>>', self.display_file)
        self.files_listbox.bind('<Double-Button-1>', self.open_file)

        # 图像显示区域
        self.image_label = ttk.Label(files_display_frame)
        self.image_label.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # 加载历史URL
        self.load_url_history()

        # === 内容选择区域 ===
        content_selection_frame = ttk.LabelFrame(control_frame, text="内容选择", padding="5")
        content_selection_frame.pack(fill=tk.X, pady=5)

        # CSS选择器
        self.enable_css_selector_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(content_selection_frame, text="启用CSS选择器",
                        variable=self.enable_css_selector_var).pack(anchor=tk.W)
        self.css_selector_text = scrolledtext.ScrolledText(
            content_selection_frame, height=3, state=tk.DISABLED)
        self.css_selector_text.pack(fill=tk.X, pady=2)
        
        self.enable_css_selector_var.trace('w', self.toggle_css_selector)

        # 标签排除
        self.enable_tag_exclusion_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(content_selection_frame, text="启用标签排除",
                        variable=self.enable_tag_exclusion_var).pack(anchor=tk.W)
        self.excluded_tags_text = scrolledtext.ScrolledText(
            content_selection_frame, height=3, state=tk.DISABLED)
        self.excluded_tags_text.pack(fill=tk.X, pady=2)
        
        self.enable_tag_exclusion_var.trace('w', self.toggle_tag_exclusion)

        # 域名排除
        self.enable_domain_exclusion_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(content_selection_frame, text="启用域名排除",
                        variable=self.enable_domain_exclusion_var).pack(anchor=tk.W)
        self.excluded_domains_text = scrolledtext.ScrolledText(
            content_selection_frame, height=3, state=tk.DISABLED)
        self.excluded_domains_text.pack(fill=tk.X, pady=2)
        
        self.enable_domain_exclusion_var.trace('w', self.toggle_domain_exclusion)

        # 社交媒体链接
        self.exclude_social_media_links_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(content_selection_frame, text="排除社交媒体链接",
                        variable=self.exclude_social_media_links_var).pack(anchor=tk.W)

        # === iframe处理区域 ===
        iframe_frame = ttk.LabelFrame(content_selection_frame, text="iframe处理", padding="5")
        iframe_frame.pack(fill=tk.X, pady=5)

        self.process_iframes_var = tk.BooleanVar(value=False)
        self.remove_overlay_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(iframe_frame, text="处理iframe内容",
                        variable=self.process_iframes_var).pack(anchor=tk.W)
        ttk.Checkbutton(iframe_frame, text="移除遮罩层",
                        variable=self.remove_overlay_var).pack(anchor=tk.W)

        # === 媒体选择区域 ===
        media_selection_frame = ttk.LabelFrame(content_selection_frame, text="媒体选择", padding="5")
        media_selection_frame.pack(fill=tk.X, pady=5)

        # 媒体类型选择
        self.enable_media_filter_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(media_selection_frame, text="启用媒体过滤",
                        variable=self.enable_media_filter_var).pack(anchor=tk.W)

        media_types_frame = ttk.Frame(media_selection_frame)
        media_types_frame.pack(fill=tk.X, pady=2)
        
        self.include_images_var = tk.BooleanVar(value=True)
        self.include_videos_var = tk.BooleanVar(value=True)
        self.include_audios_var = tk.BooleanVar(value=True)
        
        ttk.Checkbutton(media_types_frame, text="包含图片",
                        variable=self.include_images_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(media_types_frame, text="包含视频",
                        variable=self.include_videos_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(media_types_frame, text="包含音频",
                        variable=self.include_audios_var).pack(side=tk.LEFT, padx=5)

        # 媒体数据过滤
        metadata_frame = ttk.Frame(media_selection_frame)
        metadata_frame.pack(fill=tk.X, pady=2)
        
        ttk.Label(metadata_frame, text="最低相关度分数:").pack(side=tk.LEFT)
        self.media_score_threshold_var = tk.DoubleVar(value=0.0)
        ttk.Entry(metadata_frame, textvariable=self.media_score_threshold_var,
                  width=10).pack(side=tk.LEFT, padx=5)

        # === 内容处理区域 ===
        content_processing_frame = ttk.LabelFrame(control_frame, text="内容处理", padding="5")
        content_processing_frame.pack(fill=tk.X, pady=5)

        # 基本清理选项
        basic_cleaning_frame = ttk.Frame(content_processing_frame)
        basic_cleaning_frame.pack(fill=tk.X, pady=2)

        self.remove_noise_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(basic_cleaning_frame, text="移除噪音内容",
                        variable=self.remove_noise_var).pack(side=tk.LEFT, padx=5)

        self.smart_extract_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(basic_cleaning_frame, text="智能内容提取",
                        variable=self.smart_extract_var).pack(side=tk.LEFT, padx=5)

        # 内容相关性设置
        relevance_frame = ttk.Frame(content_processing_frame)
        relevance_frame.pack(fill=tk.X, pady=2)

        ttk.Label(relevance_frame, text="内容相关度阈值:").pack(side=tk.LEFT)
        self.content_relevance_var = tk.DoubleVar(value=0.5)
        ttk.Entry(relevance_frame, textvariable=self.content_relevance_var,
                  width=10).pack(side=tk.LEFT, padx=5)

        # 内容过滤选项
        filter_frame = ttk.LabelFrame(content_processing_frame, text="内容过滤", padding="5")
        filter_frame.pack(fill=tk.X, pady=2)

        self.filter_options = {
            'remove_ads': tk.BooleanVar(value=True),
            'remove_social': tk.BooleanVar(value=True),
            'remove_navigation': tk.BooleanVar(value=True),
            'remove_sidebars': tk.BooleanVar(value=True),
            'remove_footers': tk.BooleanVar(value=True)
        }

        filter_options_frame = ttk.Frame(filter_frame)
        filter_options_frame.pack(fill=tk.X)

        # 创建两列布局
        left_filters = ttk.Frame(filter_options_frame)
        left_filters.pack(side=tk.LEFT, fill=tk.X, expand=True)
        right_filters = ttk.Frame(filter_options_frame)
        right_filters.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # 左侧过滤选项
        ttk.Checkbutton(left_filters, text="移除广告", 
                        variable=self.filter_options['remove_ads']).pack(anchor=tk.W)
        ttk.Checkbutton(left_filters, text="移除社交元素", 
                        variable=self.filter_options['remove_social']).pack(anchor=tk.W)
        ttk.Checkbutton(left_filters, text="移除导航", 
                        variable=self.filter_options['remove_navigation']).pack(anchor=tk.W)

        # 右侧过滤选项
        ttk.Checkbutton(right_filters, text="移除侧边栏", 
                        variable=self.filter_options['remove_sidebars']).pack(anchor=tk.W)
        ttk.Checkbutton(right_filters, text="移除页脚", 
                        variable=self.filter_options['remove_footers']).pack(anchor=tk.W)

        # 元数据提取选项
        metadata_frame = ttk.LabelFrame(content_processing_frame, text="元数据提取", padding="5")
        metadata_frame.pack(fill=tk.X, pady=2)

        self.extract_metadata_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(metadata_frame, text="提取元数据",
                        variable=self.extract_metadata_var,
                        command=self.toggle_metadata_options).pack(anchor=tk.W)

        self.metadata_options_frame = ttk.Frame(metadata_frame)
        self.metadata_options_frame.pack(fill=tk.X)

        # 创建两列布局
        left_metadata = ttk.Frame(self.metadata_options_frame)
        left_metadata.pack(side=tk.LEFT, fill=tk.X, expand=True)
        right_metadata = ttk.Frame(self.metadata_options_frame)
        right_metadata.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # 元数据选项
        self.metadata_options = {
            'extract_title': tk.BooleanVar(value=True),
            'extract_description': tk.BooleanVar(value=True),
            'extract_keywords': tk.BooleanVar(value=True),
            'extract_author': tk.BooleanVar(value=True),
            'extract_dates': tk.BooleanVar(value=True),
            'extract_language': tk.BooleanVar(value=True),
            'extract_readability': tk.BooleanVar(value=True),
            'extract_text_stats': tk.BooleanVar(value=True)
        }

        # 左侧选项
        ttk.Checkbutton(left_metadata, text="标题", 
                        variable=self.metadata_options['extract_title']).pack(anchor=tk.W)
        ttk.Checkbutton(left_metadata, text="描述", 
                        variable=self.metadata_options['extract_description']).pack(anchor=tk.W)
        ttk.Checkbutton(left_metadata, text="关键词", 
                        variable=self.metadata_options['extract_keywords']).pack(anchor=tk.W)
        ttk.Checkbutton(left_metadata, text="可读性分析", 
                        variable=self.metadata_options['extract_readability']).pack(anchor=tk.W)

        # 右侧选项
        ttk.Checkbutton(right_metadata, text="作者", 
                        variable=self.metadata_options['extract_author']).pack(anchor=tk.W)
        ttk.Checkbutton(right_metadata, text="日期", 
                        variable=self.metadata_options['extract_dates']).pack(anchor=tk.W)
        ttk.Checkbutton(right_metadata, text="语言", 
                        variable=self.metadata_options['extract_language']).pack(anchor=tk.W)
        ttk.Checkbutton(right_metadata, text="文本统计", 
                        variable=self.metadata_options['extract_text_stats']).pack(anchor=tk.W)

        # 高级内容处理选项
        advanced_content_frame = ttk.LabelFrame(content_processing_frame, text="高级内容处理", padding="5")
        advanced_content_frame.pack(fill=tk.X, pady=2)

        self.enable_content_analysis_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(advanced_content_frame, text="启用内容分析",
                        variable=self.enable_content_analysis_var,
                        command=self.toggle_content_analysis).pack(anchor=tk.W)

        self.content_analysis_frame = ttk.Frame(advanced_content_frame)
        self.content_analysis_frame.pack(fill=tk.X)
        # 内容分析选项
        self.content_analysis_options = {
            'sentiment_analysis': tk.BooleanVar(value=False),
            'topic_detection': tk.BooleanVar(value=False),
            'entity_recognition': tk.BooleanVar(value=False),
            'summary_generation': tk.BooleanVar(value=False)
        }

        for text, var in [
            ("情感分析", 'sentiment_analysis'),
            ("题检", 'topic_detection'),
            ("实体识别", 'entity_recognition'),
            ("摘要生成", 'summary_generation')
        ]:
            ttk.Checkbutton(self.content_analysis_frame, text=text,
                            variable=self.content_analysis_options[var]).pack(anchor=tk.W)

        # 初始状态设置
        self.toggle_metadata_options()
        self.toggle_content_analysis()

        # 添加内容处理器存储
        self._content_processors = {
            'text': self.process_text_content,
            'markdown': lambda c, o: self.process_markdown_content(c, o, False),
            'fit_markdown': lambda c, o: self.process_markdown_content(c, o, True),
            'html': lambda c, o: self.process_html_content(c, o, False),
            'cleaned_html': lambda c, o: self.process_html_content(c, o, True)
        }
        
        # 添加结果缓存
        self._last_result = None
        
        # 网页克隆选项
        clone_frame = ttk.Frame(browser_frame)
        clone_frame.pack(fill=tk.X, pady=5)
        self.enable_page_clone = tk.BooleanVar(value=False)
        ttk.Checkbutton(clone_frame, text="启用网页克隆",
                        variable=self.enable_page_clone).pack(side=tk.LEFT, padx=5)
        
        # 添加进度条
        self.progress_frame = ttk.Frame(browser_frame)
        self.progress_var = tk.DoubleVar(value=0.0)
        self.progress_bar = ttk.Progressbar(
            self.progress_frame, 
            variable=self.progress_var,
            maximum=100,
            mode='determinate'
        )
        self.progress_label = ttk.Label(self.progress_frame, text="")
        
        # 进度条和标签默认隐藏
        self.progress_frame.pack_forget()

        # 纯文本提取选项
        text_frame = ttk.Frame(browser_frame)
        text_frame.pack(fill=tk.X, pady=5)
        
        # 创建水平框架来放置两个选项
        text_options_row = ttk.Frame(text_frame)
        text_options_row.pack(fill=tk.X)
        
        # 纯文本提取开关
        text_extract_frame = ttk.Frame(text_options_row)
        text_extract_frame.pack(side=tk.LEFT, padx=(0, 10))
        self.enable_text_extract = tk.BooleanVar(value=False)
        ttk.Checkbutton(text_extract_frame, text="启用纯文本提取",
                        variable=self.enable_text_extract,
                        command=self.toggle_llm_optimize).pack(side=tk.LEFT)

        # LLM文本优化开关
        llm_optimize_frame = ttk.Frame(text_options_row)
        llm_optimize_frame.pack(side=tk.LEFT)
        self.enable_llm_optimize = tk.BooleanVar(value=False)
        self.llm_optimize_cb = ttk.Checkbutton(llm_optimize_frame, text="启用LLM文本优化",
                                              variable=self.enable_llm_optimize,
                                              command=self.toggle_llm_settings,
                                              state=tk.DISABLED)
        self.llm_optimize_cb.pack(side=tk.LEFT)

        # LLM设置框架
        self.llm_settings_frame = ttk.LabelFrame(text_frame, text="LLM优化设置")
        
        # 模型类型选择 (只保留这一处定义)
        model_type_frame = ttk.Frame(self.llm_settings_frame)
        model_type_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(model_type_frame, text="选择模型类型:").pack(side=tk.LEFT)
        
        self.model_type_var = tk.StringVar(value="local")  # 设置初始值
        ttk.Radiobutton(model_type_frame, text="本地模型", 
                        variable=self.model_type_var, 
                        value="local",
                        command=self.toggle_model_settings).pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(model_type_frame, text="API模型", 
                        variable=self.model_type_var, 
                        value="api",
                        command=self.toggle_model_settings).pack(side=tk.LEFT, padx=5)

        # 本地模型设置���架 (只保留这一处定义)
        self.local_model_frame = ttk.LabelFrame(self.llm_settings_frame, text="本地模型设置")
        self.local_model_frame.pack(fill=tk.X, padx=5, pady=5)

        # 本地模型选择
        llm_model_frame = ttk.Frame(self.local_model_frame)
        llm_model_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(llm_model_frame, text="选择模型:").pack(side=tk.LEFT)

        self.llm_model_var = tk.StringVar()
        self.llm_model_combo = ttk.Combobox(llm_model_frame, 
                                           textvariable=self.llm_model_var,
                                           values=self.scan_local_models())
        self.llm_model_combo.pack(side=tk.LEFT, padx=5)

        # 添加模型管理按钮
        ttk.Button(llm_model_frame, text="模型管理", 
                  command=self.show_model_manager).pack(side=tk.LEFT, padx=5)

        # 刷新模型列表按钮
        ttk.Button(llm_model_frame, text="刷新列表", 
                  command=lambda: self.refresh_models("local")).pack(side=tk.LEFT, padx=5)

        # API模型设置框架 (只保留这一处定义)
        self.api_model_frame = ttk.LabelFrame(self.llm_settings_frame, text="API模型设置")
        self.api_model_frame.pack(fill=tk.X, padx=5, pady=5)
        self.api_model_frame.pack_forget()  # 默认隐藏

        # API提供商选择
        api_provider_frame = ttk.Frame(self.api_model_frame)
        api_provider_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(api_provider_frame, text="API提供商:").pack(side=tk.LEFT)

        self.api_providers = {
            "Gitee AI": {
                "base_url": "https://ai.gitee.com/v1",
                "models": [
                    "Qwen2.5-72B-Instruct",
                    "Qwen2.5-32B-Instruct", 
                    "Qwen2.5-14B-Instruct",
                    "Qwen2.5-7B-Instruct",
                    "Qwen2.5-Coder-32B-Instruct",
                    "Yi-34B-Chat",
                    "deepseek-coder-33B-instruct",
                    "glm-4-9b-chat",
                    "Qwen2-72B-Instruct",
                    "Qwen2-7B-Instruct",
                    "code-raccoon-v1",
                    "codegeex4-all-9b"
                ]
            },
            "OpenAI": {
                "base_url": "https://api.openai.com/v1",
                "models": ["gpt-4", "gpt-3.5-turbo"]
            },
            "自定义": {
                "base_url": "",
                "models": []
            }
        }

        self.api_provider_var = tk.StringVar(value="Gitee AI")
        self.api_provider_combo = ttk.Combobox(api_provider_frame, 
                                          textvariable=self.api_provider_var,
                                          values=list(self.api_providers.keys()))
        self.api_provider_combo.pack(side=tk.LEFT, padx=5)
        self.api_provider_combo.bind('<<ComboboxSelected>>', self.on_api_provider_change)

        # API URL输入框
        api_url_frame = ttk.Frame(self.api_model_frame)
        api_url_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(api_url_frame, text="API URL:").pack(side=tk.LEFT)
        self.api_url_var = tk.StringVar(value=self.api_providers["Gitee AI"]["base_url"])
        self.api_url_entry = ttk.Entry(api_url_frame, textvariable=self.api_url_var)
        self.api_url_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        # 修改 API Key 输入框部分的布局
        api_key_frame = ttk.Frame(self.api_model_frame)
        api_key_frame.pack(fill=tk.X, padx=5, pady=5)

        # 创建左侧框架用于标签和输入框
        left_frame = ttk.Frame(api_key_frame)
        left_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)

        ttk.Label(left_frame, text="API Key:").pack(side=tk.LEFT)

        # 修改默认值为两个 key 的组合
        default_keys = [
            "99ZE2NVXCNLWIVWC6HQBGV5GMIKCEA9D8FXL16XN",
            "R6XZ3CRX2ZXWZ5XLCR3CLHDRNNQB6OAHYHTMJCU6"  # 更新为新的 Key
        ]
        self.api_key_var = tk.StringVar(value=",".join(default_keys))
        self.api_key_entry = ttk.Entry(left_frame, textvariable=self.api_key_var, show="*")
        self.api_key_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        # 创建右侧框架用于按钮和指示器
        right_frame = ttk.Frame(api_key_frame)
        right_frame.pack(side=tk.RIGHT)

        # 添加当前使用的 Key 标签
        self.key_indicator = ttk.Label(right_frame, text="使用: Key 1/2")
        self.key_indicator.pack(side=tk.LEFT, padx=5)

        # 添加切换 API Key 的按钮
        ttk.Button(right_frame, text="切换Key", 
          command=self.switch_api_key).pack(side=tk.LEFT, padx=5)

        # API模型选择
        api_model_select_frame = ttk.Frame(self.api_model_frame)
        api_model_select_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(api_model_select_frame, text="选择模型:").pack(side=tk.LEFT)
        self.api_model_var = tk.StringVar()
        self.api_model_combo = ttk.Combobox(api_model_select_frame, 
                                       textvariable=self.api_model_var,
                                       values=self.api_providers["Gitee AI"]["models"])
        self.api_model_combo.pack(side=tk.LEFT, padx=5)

        # 刷新API模型列表按钮
        ttk.Button(api_model_select_frame, text="刷新模型列表", 
                    command=lambda: self.refresh_models("api")).pack(side=tk.LEFT, padx=5)

        # 添加提示词设置框架
        prompt_settings_frame = ttk.LabelFrame(self.api_model_frame, text="提示词设置")
        prompt_settings_frame.pack(fill=tk.X, padx=5, pady=5)

        # 添加默认提示词
        self.default_system_prompt = """你是一个专业的文本优化助手。请按照以下要求优化文本:
1. 保持原文的核心内容和主要信息不变
2. 优化文本结构, 使段落组织更加合理
3. 改善表达方式, 使语言更加流畅自然
4. 纠正语法错误和不准确的表述
5. 调整格式, 使文本更易阅读

请直接返回优化后的文本, 无需解释修改内容."""

        # 创建提示词编辑区域
        self.enable_custom_system_prompt = tk.BooleanVar(value=False)
        ttk.Checkbutton(prompt_settings_frame, text="自定义系统提示词",
                        variable=self.enable_custom_system_prompt,
                        command=self.toggle_system_prompt).pack(anchor=tk.W)

        self.system_prompt_text = scrolledtext.ScrolledText(
            prompt_settings_frame, height=8, width=50, state=tk.DISABLED)
        self.system_prompt_text.pack(fill=tk.X, pady=2)

        # 显示默认提示词（灰色）
        self.system_prompt_text.configure(state=tk.NORMAL)
        self.system_prompt_text.insert(tk.END, self.default_system_prompt)
        self.system_prompt_text.configure(state=tk.DISABLED, fg='gray')

        # 添加重置按钮
        ttk.Button(prompt_settings_frame, text="重置为默认提示词",
                   command=self.reset_system_prompt).pack(anchor=tk.E, pady=2)

        # API参数设置框架
        api_params_frame = ttk.LabelFrame(self.api_model_frame, text="API参数设置")
        api_params_frame.pack(fill=tk.X, padx=5, pady=5)

        # API参数变量
        self.api_params = {
            'stream': tk.BooleanVar(value=False),
            'max_tokens': tk.IntVar(value=512),
            'temperature': tk.DoubleVar(value=0.7),
            'top_p': tk.DoubleVar(value=0.7),
            'top_k': tk.IntVar(value=50),
            'frequency_penalty': tk.IntVar(value=1)
        }

        # 创建参数输入界面
        def create_param_entry(parent, label, var, default_value, tooltip):
            frame = ttk.Frame(parent)
            frame.pack(fill=tk.X, pady=2)
            ttk.Label(frame, text=f"{label}:").pack(side=tk.LEFT)
            
            if isinstance(var, tk.BooleanVar):
                ttk.Checkbutton(frame, variable=var).pack(side=tk.LEFT, padx=5)
            else:
                ttk.Entry(frame, textvariable=var, width=10).pack(side=tk.LEFT, padx=5)
            
            ttk.Label(frame, text=f"(默认: {default_value})", 
                     foreground="gray").pack(side=tk.LEFT)
            ttk.Label(frame, text=f"- {tooltip}", 
                     foreground="gray").pack(side=tk.LEFT, padx=5)

        # 参数说明
        params_info = {
            'stream': ("流式响应", "False", "启用流式返回生成的文本"),
            'max_tokens': ("最大生成长度", "512", "生成文本的最大token数量"),
            'temperature': ("温度系数", "0.7", "控制生成文本的随机性，值越大越随机"),
            'top_p': ("Top P", "0.7", "控制生成文本的多��性"),
            'top_k': ("Top K", "50", "从概率最大的K个词中采样"),
            'frequency_penalty': ("频率惩罚", "1", "降低重复文本的生成概率")
        }

        # 创建参数输入界面
        for param, var in self.api_params.items():
            label, default, tooltip = params_info[param]
            create_param_entry(api_params_frame, label, var, default, tooltip)

        # 添加重置按钮
        def reset_params():
            self.api_params['stream'].set(False)
            self.api_params['max_tokens'].set(512)
            self.api_params['temperature'].set(0.7)
            self.api_params['top_p'].set(0.7)
            self.api_params['top_k'].set(50)
            self.api_params['frequency_penalty'].set(1)

        reset_frame = ttk.Frame(api_params_frame)
        reset_frame.pack(fill=tk.X, pady=5)
        ttk.Button(reset_frame, text="重置为默认值", 
                  command=reset_params).pack(side=tk.RIGHT)

        # 优化选项框架 (共用，只保留这一处定义)
        optimize_options_frame = ttk.LabelFrame(self.llm_settings_frame, text="优化选项")
        optimize_options_frame.pack(fill=tk.X, padx=5, pady=5)

        # 定义一次优化选项
        self.llm_optimize_options = {
            'improve_readability': tk.BooleanVar(value=True),
            'enhance_structure': tk.BooleanVar(value=True),
            'fix_grammar': tk.BooleanVar(value=True),
            'summarize': tk.BooleanVar(value=False),
            'translate': tk.BooleanVar(value=False)
        }

        # 创建优化选项界面
        for i, (option, var) in enumerate(self.llm_optimize_options.items()):
            row = i // 2
            col = i % 2
            option_frame = ttk.Frame(optimize_options_frame)
            option_frame.grid(row=row, column=col, padx=5, pady=2, sticky="w")
            
            text_map = {
                'improve_readability': "提升可读性",
                'enhance_structure': "优化结构",
                'fix_grammar': "修正语法",
                'summarize': "生成摘要",
                'translate': "翻译"
            }
            
            ttk.Checkbutton(option_frame, text=text_map[option],
                           variable=var).pack(side=tk.LEFT)

        # 自定义提示词框架 (共用)
        prompt_frame = ttk.LabelFrame(self.llm_settings_frame, text="自定义提示词")
        prompt_frame.pack(fill=tk.X, padx=5, pady=5)

        # 添加自定义提示词开关
        self.enable_custom_prompt = tk.BooleanVar(value=False)
        ttk.Checkbutton(prompt_frame, text="启用自定义提示词",
                        variable=self.enable_custom_prompt,
                        command=self.toggle_custom_prompt).pack(anchor=tk.W)

        # 添加提示词输入框
        self.custom_prompt_text = scrolledtext.ScrolledText(
            prompt_frame, height=4, state=tk.DISABLED)
        self.custom_prompt_text.pack(fill=tk.X, pady=2)

        # 添加说明标签
        ttk.Label(prompt_frame, text="不输入则使用默认提示词", 
                 foreground="gray").pack(anchor=tk.W)

        # 批量处理按钮 (共用)
        batch_frame = ttk.Frame(self.llm_settings_frame)
        batch_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(batch_frame, text="选择文件进行批量优化",
                   command=self.batch_optimize_files).pack(side=tk.LEFT)

        # 初始状态下隐藏LLM设置框架
        self.llm_settings_frame.pack_forget()

        # 纯文本提取选项初始化
        self.text_extract_options = {
            # 基本设置
            'remove_ads': tk.BooleanVar(value=True),
            'remove_menus': tk.BooleanVar(value=True),
            'remove_headers': tk.BooleanVar(value=True),
            'remove_footers': tk.BooleanVar(value=True),
            'remove_comments': tk.BooleanVar(value=True),
            'remove_social': tk.BooleanVar(value=True),
            
            # 内��处理
            'keep_main_content': tk.BooleanVar(value=True),
            'keep_images': tk.BooleanVar(value=False),
            'keep_tables': tk.BooleanVar(value=False),
            'keep_links': tk.BooleanVar(value=False),
            'keep_lists': tk.BooleanVar(value=True),
            'keep_formatting': tk.BooleanVar(value=True),
            
            # 文本优化
            'merge_spaces': tk.BooleanVar(value=True),
            'smart_paragraphs': tk.BooleanVar(value=True),
            'normalize_spaces': tk.BooleanVar(value=True),
            'fix_punctuation': tk.BooleanVar(value=True),
            'remove_empty_lines': tk.BooleanVar(value=True),
            'combine_short_lines': tk.BooleanVar(value=True),
            
            # 格式选项
            'save_as_word': tk.BooleanVar(value=True),
            'add_toc': tk.BooleanVar(value=True),
            'add_page_numbers': tk.BooleanVar(value=True),
            'add_header_footer': tk.BooleanVar(value=False),
            'use_styles': tk.BooleanVar(value=True),
            
            # 高级设置
            'extract_article': tk.BooleanVar(value=True),
            'extract_title': tk.BooleanVar(value=True),
            'extract_metadata': tk.BooleanVar(value=True),
            'clean_boilerplate': tk.BooleanVar(value=True),
            'detect_language': tk.BooleanVar(value=True),
            
            # 数值参数
            'min_text_length': tk.IntVar(value=20),
            'max_title_length': tk.IntVar(value=200),
            'paragraph_threshold': tk.IntVar(value=100),
            'image_min_size': tk.IntVar(value=100),
            'max_line_length': tk.IntVar(value=80)
        }

        # 创建选项卡式布局
        text_options_notebook = ttk.Notebook(text_frame)
        text_options_notebook.pack(fill=tk.BOTH, expand=True, pady=5, padx=5)

        # 在 create_scrollable_frame 函数定义之前添加
        # 定义选项卡内容
        tabs = {
            "基本设置": [
                ("移除广告", 'remove_ads', "移除网页中的广告内容"),
                ("移除菜单", 'remove_menus', "移除导航菜单"),
                ("移除页眉", 'remove_headers', "移除页面顶部内容"),
                ("移除页脚", 'remove_footers', "移除页面底部内容"),
                ("移除评论", 'remove_comments', "移除用户评论区"),
                ("移除社交按钮", 'remove_social', "移除社交媒体分享按钮")
            ],
            "内容处理": [
                ("保留主要内容", 'keep_main_content', "���留页面主要内容区域"),
                ("保留图片", 'keep_images', "保留文章中的图片"),
                ("保留表格", 'keep_tables', "保留数据表格"),
                ("留链接", 'keep_links', "保留超链接"),
                ("保留列表", 'keep_lists', "保留序和无序列表"),
                ("保留格式", 'keep_formatting', "保留文本格式化")
            ],
            "文本优化": [
                ("合并空白", 'merge_spaces', "合并多余的空白字符"),
                ("智能段落", 'smart_paragraphs', "智能识别段落结构"),
                ("规范化空格", 'normalize_spaces', "统一空格使用"),
                ("修正标点", 'fix_punctuation', "修正标点符号使用"),
                ("移除多余行", 'remove_empty_lines', "移除多余的空行"),
                ("合并短行", 'combine_short_lines', "合并过短的文本行")
            ],
            "格式选项": [
                ("保存为Word", 'save_as_word', "将内容保存为Word文档"),
                ("添加目录", 'add_toc', "在文档中添加目录"),
                ("添加页码", 'add_page_numbers', "添加页码"),
                ("加页眉页脚", 'add_header_footer', "添加页眉和页脚"),
                ("使用样式", 'use_styles', "应用预定义的样式")
            ],
            "高级设置": [
                ("提取文章", 'extract_article', "智能提取主要文章内容"),
                ("提取标题", 'extract_title', "提取页面标题"),
                ("提取元数据", 'extract_metadata', "提取页面元数据"),
                ("清理样板文本", 'clean_boilerplate', "移除重复的样板文本"),
                ("检测语言", 'detect_language', "检测文本语言")
            ],
            "数值参数": [
                ("最小文本长度", 'min_text_length', "设置最小文本长度阈值"),
                ("最大标题长度", 'max_title_length', "设置最大标题长度"),
                ("段落阈值", 'paragraph_threshold', "设置段落字数阈值"),
                ("最小图片尺寸", 'image_min_size', "设置��小图片尺寸"),
                ("最大行长度", 'max_line_length', "设置最大行字符数")
            ]
        }

        # 修改 create_scrollable_frame 函数
        def create_scrollable_frame(parent):
            frame = ttk.Frame(parent)
            
            # 创建水平和垂直滚动条
            h_scrollbar = ttk.Scrollbar(frame, orient="horizontal")
            v_scrollbar = ttk.Scrollbar(frame, orient="vertical")
            canvas = tk.Canvas(frame, height=150,  # 减小高度使界面更紧凑
                              xscrollcommand=h_scrollbar.set,
                              yscrollcommand=v_scrollbar.set)
            
            scrollable_frame = ttk.Frame(canvas)
            
            # 配置滚动条
            h_scrollbar.config(command=canvas.xview)
            v_scrollbar.config(command=canvas.yview)
            
            # 绑定框架大小变化事件
            def configure_scroll_region(event):
                canvas.configure(scrollregion=canvas.bbox("all"))
            scrollable_frame.bind("<Configure>", configure_scroll_region)
            
            # 创建窗口
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            
            # 绑定鼠标滚轮事件
            def _on_mousewheel(event):
                if event.state == 0:  # 没有按住Shift键
                    canvas.yview_scroll(int(-1*(event.delta/120)), "units")
                else:  # 按住Shift键进行水平滚动
                    canvas.xview_scroll(int(-1*(event.delta/120)), "units")
            
            canvas.bind_all("<MouseWheel>", _on_mousewheel)
            canvas.bind_all("<Shift-MouseWheel>", _on_mousewheel)
            
            # 布局
            canvas.grid(row=0, column=0, sticky="nsew")
            v_scrollbar.grid(row=0, column=1, sticky="ns")
            h_scrollbar.grid(row=1, column=0, sticky="ew")
            
            # 配置网格权重
            frame.grid_rowconfigure(0, weight=1)
            frame.grid_columnconfigure(0, weight=1)
            
            return frame, scrollable_frame

        # 修改选项卡内容的创建部分
        for tab_name, options in tabs.items():
            frame, scrollable_frame = create_scrollable_frame(text_options_notebook)
            text_options_notebook.add(frame, text=tab_name)
            
            # 使用网格布局来排列选项
            row = 0
            col = 0
            max_cols = 2  # 每行显示的选项数
            
            for i, (text, option_name, tooltip) in enumerate(options):
                option_frame = ttk.Frame(scrollable_frame)
                option_frame.grid(row=row, column=col, padx=5, pady=2, sticky="w")
                
                if tab_name != "数值参数":
                    # 复选框选项
                    var = self.text_extract_options[option_name]
                    cb = ttk.Checkbutton(option_frame, text=text, variable=var)
                    cb.pack(side=tk.LEFT)
                    
                    # 添加工具示标签（使用较短的提示文本）
                    ttk.Label(option_frame, text=f"({tooltip})", 
                             foreground="gray", wraplength=200).pack(side=tk.LEFT, padx=5)
                else:
                    # 数值输入��项
                    ttk.Label(option_frame, text=text).pack(side=tk.LEFT)
                    ttk.Entry(option_frame, textvariable=self.text_extract_options[option_name], 
                             width=6).pack(side=tk.LEFT, padx=2)
                    ttk.Label(option_frame, text=f"({tooltip})", 
                             foreground="gray", wraplength=150).pack(side=tk.LEFT, padx=2)
                
                # 更新行列位置
                col += 1
                if col >= max_cols:
                    col = 0
                    row += 1

        # 设置选项卡整体高度和宽度
        text_options_notebook.configure(height=180, width=400)  # 减小高度，设置合适的宽度

    def toggle_js_code(self, *args):
        if self.enable_js_code_var.get():
            self.js_code_text.config(state=tk.NORMAL)
        else:
            self.js_code_text.delete('1.0', tk.END)
            self.js_code_text.config(state=tk.DISABLED)

    def toggle_wait_for(self, *args):
        if self.enable_wait_for_var.get():
            self.wait_for_text.config(state=tk.NORMAL)
        else:
            self.wait_for_text.delete('1.0', tk.END)
            self.wait_for_text.config(state=tk.DISABLED)

    def toggle_session_id(self, *args):
        if self.enable_session_var.get():
            self.session_id_entry.config(state=tk.NORMAL)
        else:
            self.session_id_var.set("")
            self.session_id_entry.config(state=tk.DISABLED)

    def toggle_extraction_options(self):
        strategy = self.extraction_strategy_var.get()
        if strategy == "jsoncss":
            self.jsoncss_config_frame.pack(fill=tk.X, pady=2)
            self.llm_config_frame.pack_forget()
        elif strategy == "llm":
            self.llm_config_frame.pack(fill=tk.X, pady=2)
            self.jsoncss_config_frame.pack_forget()
        else:
            self.jsoncss_config_frame.pack_forget()
            self.llm_config_frame.pack_forget()

    def ensure_directories(self):
        """确保所有必要的目录都存在"""
        # 创建主要目录
        directories = {
            'data': self.base_dir / "data",
            'content': self.base_dir / "content",
            'screenshots': self.base_dir / "screenshots",
            'media': self.base_dir / "media",
            'links': self.base_dir / "links"
        }

        for dir_path in directories.values():
            dir_path.mkdir(parents=True, exist_ok=True)

        self.directories = directories

    def get_safe_filename(self, url):
        """生成安全的文件名"""
        # 从URL中提取域名
        domain = urlparse(url).netloc
        if not domain:
            domain = 'unknown_domain'
        # 移除非法字符
        domain = re.sub(r'[<>:"/\\|?*]', '_', domain)
        # 生成时间戳
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # 组合文件名
        return f"{domain}_{timestamp}"

    def save_content(self, content, url, format_type):
        """保存内容到文件"""
        filename = self.get_safe_filename(url)
        file_path = self.directories['content'] / f"{filename}_{format_type}.txt"

        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            self.saved_files.append(file_path)
            return file_path
        except Exception as e:
            logging.error(f"保存内容失败: {e}")
            return None

    def save_links(self, links, url):
        """保存链接信息"""
        filename = self.get_safe_filename(url)
        file_path = self.directories['links'] / f"{filename}_links.json"

        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(links, f, ensure_ascii=False, indent=2)
            self.saved_files.append(file_path)
            return file_path
        except Exception as e:
            logging.error(f"保存链接信息失败: {e}")
            return None

    def save_media_info(self, media, url):
        """保存媒体信息"""
        filename = self.get_safe_filename(url)
        file_path = self.directories['media'] / f"{filename}_media.json"

        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(media, f, ensure_ascii=False, indent=2)
            self.saved_files.append(file_path)
            return file_path
        except Exception as e:
            logging.error(f"保存媒体信息失���: {e}")
            return None

    def save_screenshot(self, screenshot_data, url):
        """保存截图"""
        if not screenshot_data:
            logging.warning("截图数据为空")
            return None

        try:
            filename = self.get_safe_filename(url)
            screenshot_path = self.directories['screenshots'] / f"{filename}.png"

            # 检查数据类型并相应处理
            if isinstance(screenshot_data, str):
                # 如果是base64字串
                try:
                    decoded_data = base64.b64decode(screenshot_data)
                    with open(screenshot_path, "wb") as f:
                        f.write(decoded_data)
                except Exception as e:
                    logging.error(f"Base64解码失败: {e}")
                    return None
            elif isinstance(screenshot_data, bytes):
                # 如果是字节数据
                with open(screenshot_path, "wb") as f:
                    f.write(screenshot_data)
            else:
                logging.error(f"不支持的截图数据类型: {type(screenshot_data)}")
                return None

            self.saved_files.append(screenshot_path)
            return screenshot_path
        except Exception as e:
            logging.error(f"保存截图失败: {e}")
            return None

    def load_url_history(self):
        """加载URL历史记录"""
        try:
            if self.urls_file.exists():
                with open(self.urls_file, 'r', encoding='utf-8') as f:
                    urls = json.load(f)
                self.url_combobox['values'] = urls
                if urls:
                    self.url_combobox.set(urls[0])
                else:
                    self.url_combobox.set("https://example.com")
            else:
                self.url_combobox.set("https://example.com")
        except Exception as e:
            logging.error(f"加载URL历史记录失败: {e}")
            self.url_combobox.set("https://example.com")

    def save_url_history(self):
        """保存URL历史记录"""
        try:
            current_url = self.url_var.get()
            urls = list(self.url_combobox['values'])

            # 当前URL移到最前面
            if current_url in urls:
                urls.remove(current_url)
            urls.insert(0, current_url)

            # 限制最多保存20个URL
            urls = urls[:20]

            self.url_combobox['values'] = urls

            with open(self.urls_file, 'w', encoding='utf-8') as f:
                json.dump(urls, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logging.error(f"保存URL历史记录失败: {e}")

    def delete_current_url(self):
        """删除当前选中的URL"""
        current_url = self.url_var.get()
        urls = list(self.url_combobox['values'])
        if current_url in urls:
            urls.remove(current_url)
            self.url_combobox['values'] = urls
            if urls:
                self.url_combobox.set(urls[0])
            else:
                self.url_combobox.set("https://example.com")
            self.save_url_history()

    def start_crawl(self):
        self.crawl_button.configure(state='disabled')
        # 开始新的线程来运行异步爬取，以防止阻塞Tkinter主线程
        threading.Thread(target=self.run_crawl).start()

    def run_crawl(self):
        asyncio.run(self.crawl())

    async def crawl(self):
        """优化的爬取方法"""
        error_message = None
        try:
            # 禁用爬取按钮显示状态
            self.root.after(0, lambda: (
                self.crawl_button.configure(state='disabled'),
                self.content_text.insert(tk.END, "爬取中...\n")
            ))
            
            # 构建配置
            crawler_config, crawl_config = self._build_configs()
            
            # 执行爬取
            async with AsyncWebCrawler(**crawler_config) as crawler:
                self.save_url_history()
                result = await crawler.arun(**crawl_config)
                
                # 缓存结果
                self._last_result = result
                
                # 异步处理结果
                await self._process_result(result)
                
        except Exception as e:
            error_message = str(e)
            logging.exception("爬取过程发生错误")
            self.root.after(0, lambda: self.content_text.insert(tk.END, f"错误: {error_message}\n"))
            
        finally:
            self.root.after(0, lambda: self.crawl_button.configure(state='normal'))

    async def _process_result(self, result):
        """异步处理爬取结果"""
        if not result:
            self.root.after(0, lambda: self.content_text.insert(tk.END, "未获取到结果\n"))
            return
            
        try:
            # 获取内容
            content = self._extract_content(result)
            if not content:
                raise ValueError("未能提取内容")
                
            # 处理内容
            format_type = self.output_format.get()
            processor = self._content_processors.get(format_type)
            if not processor:
                raise ValueError(f"不支持的格式类型: {format_type}")
                
            # 获取格式化选项
            format_options = {k: v.get() for k, v in self.format_options.items()}
            
            # 处理内容
            processed_content = processor(content, format_options)
            
            # 准备结果数据
            url = self.url_var.get()
            saved_files = {}
            result_data = {
                'content': processed_content,
                'media': {},
                'links': {},
                'metadata': {},
                'analysis': {}
            }
            
            try:
                # 保存主要内容
                saved_files['content'] = await self._save_content_async(
                    processed_content, url, format_type)
                
                # 处理媒体信息
                if hasattr(result, 'media') and result.media:
                    result_data['media'] = result.media
                    media_path = await self._save_json_async(
                        result.media, 
                        self.directories['media'] / f"{self.get_safe_filename(url)}_media.json")
                saved_files['media'] = media_path
                
                # 处理链接信息
                if hasattr(result, 'links') and result.links:
                    result_data['links'] = result.links
                    links_path = await self._save_json_async(
                        result.links,
                        self.directories['links'] / f"{self.get_safe_filename(url)}_links.json")
                saved_files['links'] = links_path
                
                # 处理元数据
                metadata_path = None  # 初始化为 None
                if self.extract_metadata_var.get():
                    metadata = {}
                    if hasattr(result, 'title'):
                        metadata['title'] = result.title
                    if hasattr(result, 'description'):
                        metadata['description'] = result.description
                    if hasattr(result, 'keywords'):
                        metadata['keywords'] = result.keywords
                    if hasattr(result, 'author'):
                        metadata['author'] = result.author
                    if hasattr(result, 'language'):
                        metadata['language'] = result.language
                    if hasattr(result, 'publish_date'):
                        metadata['publish_date'] = str(result.publish_date)
                    
                    if metadata:  # 只有在有元数据时才保存
                        result_data['metadata'] = metadata
                        metadata_path = await self._save_json_async(
                            metadata,
                            self.directories['data'] / f"{self.get_safe_filename(url)}_metadata.json")
                saved_files['metadata'] = metadata_path  # 无论是否有元数据，都添加到saved_files中
                
                # 处理内容分析
                analysis_path = None  # 初始化为 None
                if self.enable_content_analysis_var.get():
                    analysis = {}
                    if hasattr(result, 'sentiment'):
                        analysis['sentiment'] = result.sentiment
                    if hasattr(result, 'topics'):
                        analysis['topics'] = result.topics
                    if hasattr(result, 'entities'):
                        analysis['entities'] = result.entities
                    if hasattr(result, 'summary'):
                        analysis['summary'] = result.summary
                    
                    if analysis:  # 只有在有分析数据时才保存
                        result_data['analysis'] = analysis
                        analysis_path = await self._save_json_async(
                            analysis,
                            self.directories['data'] / f"{self.get_safe_filename(url)}_analysis.json")
                saved_files['analysis'] = analysis_path  # 无论是否有分析数据，都添加到saved_files中

                # 保存截图（如果启用）
                screenshot_path = None  # 初始化为 None
                if self.screenshot_var.get() and hasattr(result, 'screenshot'):
                    screenshot_path = await self._save_screenshot_async(
                        result.screenshot,
                        url)
                saved_files['screenshot'] = screenshot_path  # 无论是否有截图，都添加到saved_files中

                # 保存可浏览网页
                if hasattr(result, 'html'):
                    browsable_page = await self.save_browsable_page(
                        result.html,
                        self.url_var.get(),
                        getattr(result, 'resources', None))
                    if browsable_page:
                        self.root.after(0, lambda: self.content_text.insert(
                            tk.END, 
                            f"\n\n可浏览网页已保存至: {browsable_page}\n"))

                # 提取纯文本
                if hasattr(result, 'html'):
                    text_path = await self.extract_pure_text(result.html)
                    if text_path:
                        self.root.after(0, lambda: self.content_text.insert(
                            tk.END, 
                            f"\n\n纯文本已保存至: {text_path}\n"))

            except Exception as save_error:
                logging.error(f"保存文件时发生错误: {save_error}")
                
            # 更新GUI显示
            self.root.after(0, lambda: self._update_display(
                content=result_data['content'],
                media=result_data['media'],
                links=result_data['links'],
                saved_files=saved_files,
                metadata=result_data['metadata'],
                analysis=result_data['analysis']))
            
        except Exception as e:
            logging.exception("处理结果时发生错误")
            self.root.after(0, lambda: self.content_text.insert(tk.END, f"处理错误: {str(e)}\n"))

    def _extract_content(self, result):
        """从结果中提取内容"""
        for attr in ('html', 'content', 'text'):
            if hasattr(result, attr):
                return getattr(result, attr)
        return None

    async def _save_content_async(self, content, url, format_type):
        """异步保存内容"""
        try:
            filename = self.get_safe_filename(url)
            file_path = self.directories['content'] / f"{filename}_{format_type}.txt"
            
            # 使用异步文件操作
            async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
                await f.write(content)
                
            self.saved_files.append(file_path)
            return file_path
            
        except Exception as e:
            logging.error(f"存内容失败: {e}")
            return None

    def _update_display(self, content=None, media=None, links=None, saved_files=None, metadata=None, analysis=None):
        """更新显示容"""
        # 清空现有内容
        self.clear_texts()
        
        # 显示内容
        if content:
            self.content_text.insert(tk.END, content)
        
        # 显示元数据
        if metadata:
            if self.extract_metadata_var.get():
                self.content_text.insert(tk.END, "\n\n=== 元数据 ===\n")
                for key, value in metadata.items():
                    if value and self.metadata_options.get(f'extract_{key}', tk.BooleanVar(value=False)).get():
                        self.content_text.insert(tk.END, f"{key}: {value}\n")
        
        # 显示内容分析结果
        if analysis:
            if self.enable_content_analysis_var.get():
                self.content_text.insert(tk.END, "\n\n=== 内容分析 ===\n")
                for key, value in analysis.items():
                    if value and self.content_analysis_options.get(key, tk.BooleanVar(value=False)).get():
                        self.content_text.insert(tk.END, f"{key}: {value}\n")
        
        # 显示媒体信息
        if media:
            self.display_media_info(media)
        
        # 显示链接信息
        if links:
            self.display_links_info(links)
        
        # 更新文件列表
        if saved_files:
            for file_type, file_path in saved_files.items():
                if file_path:
                    self.files_listbox.insert(tk.END, str(file_path))
                    if file_type == 'content':
                        self.content_text.insert(tk.END, f"\n\n内容已保存至: {file_path}")

    def _build_configs(self):
        """构建爬取配置"""
        try:
            # 获取基本配置
            crawler_config = {
                'browser_type': self.browser_type.get(),
                'headless': self.headless_var.get(),
                'verbose': self.verbose_var.get()
            }

            # 构建爬取配置
            crawl_config = {
                'url': self.url_var.get(),
                'word_count_threshold': self.word_count_var.get(),
                'exclude_external_links': self.exclude_external_links_var.get(),
                'exclude_external_images': self.exclude_external_images_var.get(),
                'timeout': self.timeout_var.get(),
                'delay': self.delay_var.get(),
                'js_only': self.js_only_var.get(),
                'simulate_user': self.simulate_user_var.get(),
                'magic': self.magic_var.get(),
                'screenshot': self.screenshot_var.get()
            }

            # 添加内容处理配置
            content_processing_config = self.get_content_processing_config()
            crawl_config.update(content_processing_config)

            # 添加媒体过滤配置
            if self.enable_media_filter_var.get():
                crawl_config['media_filter'] = {
                    'include_images': self.include_images_var.get(),
                    'include_videos': self.include_videos_var.get(),
                    'include_audios': self.include_audios_var.get(),
                    'score_threshold': self.media_score_threshold_var.get()
                }

            # 添加JavaScript代码配置
            if self.enable_js_code_var.get():
                js_code = self.js_code_text.get('1.0', tk.END).strip()
                if js_code:
                    crawl_config['js_code'] = js_code

            # 添加等待条件配置
            if self.enable_wait_for_var.get():
                wait_for = self.wait_for_text.get('1.0', tk.END).strip()
                if wait_for:
                    crawl_config['wait_for'] = wait_for

            # 添加会话管理配置
            if self.enable_session_var.get():
                session_id = self.session_id_var.get().strip()
                if session_id:
                    crawl_config['session_id'] = session_id

            # 添加iframe处理配置
            if self.process_iframes_var.get():
                crawl_config['process_iframes'] = True
            if self.remove_overlay_var.get():
                crawl_config['remove_overlay'] = True

            # 验证配置
            if not self.validate_config(crawl_config):
                raise ValueError("配置验证失败")

            return crawler_config, crawl_config

        except Exception as e:
            logging.error(f"构建配置时发生错误: {e}")
            raise

    def validate_config(self, config):
        """验证配置是否有效"""
        try:
            # 检查必需的配置项
            required_fields = ['url', 'word_count_threshold', 'timeout']
            for field in required_fields:
                if field not in config:
                    raise ValueError(f"缺少必需的配置项: {field}")

            # 验证URL
            url = config['url']
            if not url or not url.startswith(('http://', 'https://')):
                raise ValueError("无效的URL")

            # 验证数值类型的配置
            int_fields = ['word_count_threshold', 'timeout']
            float_fields = ['delay', 'content_relevance_threshold']

            for field in int_fields:
                if field in config and not isinstance(config[field], (int, float)):
                    raise ValueError(f"{field} 必须是数字")

            for field in float_fields:
                if field in config and not isinstance(config[field], (int, float)):
                    raise ValueError(f"{field} 必须是数字")

            # 验证内容处理配置
            if 'content_relevance_threshold' in config:
                threshold = config['content_relevance_threshold']
                if not 0 <= threshold <= 1:
                    raise ValueError("内容相关度阈值必须在0到1之��")

            # 验证媒体过滤配置
            if 'media_filter' in config:
                media_filter = config['media_filter']
                if 'score_threshold' in media_filter:
                    score = media_filter['score_threshold']
                    if not 0 <= score <= 1:
                        raise ValueError("体相关度分数必须在0到1之间")

            return True

        except Exception as e:
            logging.error(f"配置验证失败: {e}")
            return False

    def process_text_content(self, content, options):
        """处理纯文本内容"""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            
            # 创建结构化文本
            structured_text = []
            
            # 处理标题
            if options['preserve_headings']:
                for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    level = int(heading.name[1])
                    prefix = '#' * level + ' '
                    structured_text.append(f"\n{prefix}{heading.get_text().strip()}\n")
            
            # 处理段落和其他内容
            for element in soup.find_all(['p', 'div', 'section']):
                # 处理图片
                if options['preserve_images']:
                    for img in element.find_all('img'):
                        alt_text = img.get('alt', '图片')
                        src = img.get('src', '')
                        structured_text.append(f"[图片: {alt_text}]({src})\n")
                
                # 处理链接
                if options['preserve_links']:
                    for link in element.find_all('a'):
                        href = link.get('href', '')
                        text = link.get_text().strip()
                        structured_text.append(f"[链接: {text}]({href})\n")
                
                # 处理表格
                if options['preserve_tables']:
                    for table in element.find_all('table'):
                        structured_text.append("\n[表格开始]\n")
                        for row in table.find_all('tr'):
                            cells = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                            structured_text.append(" | ".join(cells))
                        structured_text.append("[表格结束]\n")
                
                # 处理列表
                if options['preserve_lists']:
                    for ul in element.find_all('ul'):
                        structured_text.append("\n")
                        for li in ul.find_all('li'):
                            structured_text.append(f"• {li.get_text().strip()}\n")
                    for ol in element.find_all('ol'):
                        structured_text.append("\n")
                        for i, li in enumerate(ol.find_all('li'), 1):
                            structured_text.append(f"{i}. {li.get_text().strip()}\n")
                
                # 处理代码块
                if options['preserve_code']:
                    for code in element.find_all(['code', 'pre']):
                        structured_text.append("\n```\n")
                        structured_text.append(code.get_text().strip())
                        structured_text.append("\n```\n")
                
                # 处理引用
                if options['preserve_quotes']:
                    for quote in element.find_all('blockquote'):
                        structured_text.append("\n> ")
                        structured_text.append(quote.get_text().strip().replace('\n', '\n> '))
                        structured_text.append("\n")
                
                # 处理强调文本
                if options['preserve_emphasis']:
                    for strong in element.find_all(['strong', 'b']):
                        text = strong.get_text().strip()
                        structured_text.append(f"**{text}**")
                    for em in element.find_all(['em', 'i']):
                        text = em.get_text().strip()
                        structured_text.append(f"*{text}*")
                
                # 处理普通文本
                text = element.get_text().strip()
                if text and not any(text in s for s in structured_text):
                    structured_text.append(text + "\n")
            
            # 合并所有文本并清理
            final_text = "\n".join(structured_text)
            # 清理多余的空行
            final_text = "\n".join(line for line in final_text.split("\n") if line.strip())
            return final_text
            
        except Exception as e:
            logging.error(f"文本处理错误: {str(e)}")
            return content

    def process_markdown_content(self, content, options, is_fit=False):
        """处理Markdown内容"""
        try:
            import html2text
            h = html2text.HTML2Text()
            
            # 配置基本选项
            h.body_width = 0  # 禁自动换行
            h.unicode_snob = True  # 使用 Unicode 字符
            h.skip_internal_links = False
            h.inline_links = True
            h.wrap_links = False
            
            # 根据保留选项配置转换器
            h.ignore_images = not options['preserve_images']
            h.ignore_links = not options['preserve_links']
            h.ignore_tables = not options['preserve_tables']
            h.ignore_emphasis = not options['preserve_emphasis']
            
            # 转换为Markdown
            markdown = h.handle(content)
            
            # 根据选项进行后处理
            lines = markdown.split('\n')
            processed_lines = []
            
            for line in lines:
                # 处理标题
                if not options['preserve_headings'] and line.startswith('#'):
                    line = line.lstrip('#').strip()
                
                # 处理列表
                if not options['preserve_lists']:
                    if line.startswith('*') or line.startswith('-') or line.startswith('+'):
                        line = line.lstrip('*-+ ').strip()
                    elif line.strip().startswith('1.'):
                        line = line.split('.', 1)[1].strip()
                
                # 处理引用
                if not options['preserve_quotes'] and line.startswith('>'):
                    line = line.lstrip('> ').strip()
                
                # 处理代码块
                if not options['preserve_code'] and line.startswith('```'):
                    continue
                
                if line.strip():
                    processed_lines.append(line)
            
            # 合并处理后的行
            processed_markdown = '\n'.join(processed_lines)
            
            # 如果是精简模式，进行额外的清理
            if is_fit:
                processed_markdown = self.fit_markdown(processed_markdown)
            
            return processed_markdown
            
        except Exception as e:
            logging.error(f"Markdown处理错误: {str(e)}")
            return content

    def fit_markdown(self, markdown):
        """精简Markdown内容"""
        # 移除连续的空行
        lines = markdown.split('\n')
        processed_lines = []
        prev_empty = False
        
        for line in lines:
            if line.strip():
                processed_lines.append(line)
                prev_empty = False
            elif not prev_empty:
                processed_lines.append('')
                prev_empty = True
        
        # 移除开头和结尾的空行
        while processed_lines and not processed_lines[0].strip():
            processed_lines.pop(0)
        while processed_lines and not processed_lines[-1].strip():
            processed_lines.pop()
        
        return '\n'.join(processed_lines)

    def clear_texts(self):
        """清空所有文本区域"""
        self.content_text.delete(1.0, tk.END)
        self.media_text.delete(1.0, tk.END)
        self.links_text.delete(1.0, tk.END)
        self.files_listbox.delete(0, tk.END)
        self.image_label.config(image='')  # 清空像显示

    def update_gui(self, data):
        """更新GUI显示"""
        try:
            # 清空现有内容
            self.content_text.delete('1.0', tk.END)
            self.media_text.delete('1.0', tk.END)
            self.links_text.delete('1.0', tk.END)

            # 重新启用爬取按钮
            self.crawl_button.configure(state='normal')

            if data.get('success'):
                # 显示内容
                if data.get('content'):
                    self.content_text.insert(tk.END, f"【{data['format_type']}格式输出】\n")
                    self.content_text.insert(tk.END, data['content'])
                else:
                    self.content_text.insert(tk.END, "未获取到内容。\n")

                # 显示媒体信息
                if any(data.get('media', {}).values()):
                    self.display_media_info(data['media'])
                else:
                    self.media_text.insert(tk.END, "未找到媒体内容\n")

                # 显示链接信息
                if any(data.get('links', {}).values()):
                    self.display_links_info(data['links'])
                else:
                    self.links_text.insert(tk.END, "未找到链接\n")

                # 显示文件保存信息
                if data.get('content_file'):
                    self.content_text.insert(tk.END, f"\n\n内容已保存至: {data['content_file']}")

            elif data.get('error'):
                self.content_text.insert(tk.END, f"发生错误: {data['error']}")
            else:
                self.content_text.insert(tk.END, "未知错误: 未能获取到任何结果")

        except Exception as e:
            logging.error(f"GUI更新时发生错误: {str(e)}")
            self.content_text.insert(tk.END, f"GUI更新错误: {str(e)}")

    def display_media_info(self, media_data):
        """显示媒体信息"""
        media_info = []
        
        if media_data.get('images'):
            media_info.append("图片:")
            for img in media_data['images']:
                if isinstance(img, dict):
                    media_info.append(f"- URL: {img.get('src', 'N/A')}")
                    if img.get('alt'):
                        media_info.append(f"  描述: {img['alt']}")
                else:
                    media_info.append(f"- {img}")
            media_info.append("")

        if media_data.get('videos'):
            media_info.append("视频:")
            for video in media_data['videos']:
                if isinstance(video, dict):
                    media_info.append(f"- URL: {video.get('src', 'N/A')}")
                    if video.get('title'):
                        media_info.append(f"  标题: {video['title']}")
                else:
                    media_info.append(f"- {video}")
            media_info.append("")

        if media_info:
            self.media_text.insert(tk.END, "\n".join(media_info))
        else:
            self.media_text.insert(tk.END, "未找到媒体内容\n")

    def display_links_info(self, links_data):
        """显示链接信息"""
        if links_data.get('internal'):
            self.links_text.insert(tk.END, "内部链接:\n")
            for link in links_data['internal']:
                if isinstance(link, dict):
                    self.links_text.insert(tk.END, f"- {link.get('href', 'N/A')}\n")
                    if link.get('text'):
                        self.links_text.insert(tk.END, f"  文本: {link['text']}\n")
                else:
                    self.links_text.insert(tk.END, f"- {link}\n")

        if links_data.get('external'):
            self.links_text.insert(tk.END, "\n外部链接:\n")
            for link in links_data['external']:
                if isinstance(link, dict):
                    self.links_text.insert(tk.END, f"- {link.get('href', 'N/A')}\n")
                    if link.get('text'):
                        self.links_text.insert(tk.END, f"  文本: {link['text']}\n")
                else:
                    self.links_text.insert(tk.END, f"- {link}\n")

    def open_file(self, event):
        """打开选中的文件"""
        selection = self.files_listbox.curselection()
        if selection:
            file_path = self.files_listbox.get(selection[0])
            file_path = Path(file_path)

            try:
                if sys.platform == "win32":
                    os.startfile(file_path)
                elif sys.platform == "darwin":
                    subprocess.run(["open", file_path])
                else:
                    subprocess.run(["xdg-open", file_path])
            except Exception as e:
                logging.error(f"无法打开文件: {e}")

    def display_file(self, event):
        """在程序中显示选中的截图"""
        selection = self.files_listbox.curselection()
        if selection:
            file_path = self.files_listbox.get(selection[0])
            file_path = Path(file_path)

            # 检查文件是否为图像文件（以 .png 结尾）
            if file_path.suffix.lower() in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                try:
                    image = Image.open(file_path)
                    image.thumbnail((400, 400))  # 调整图像大小
                    self.photo = ImageTk.PhotoImage(image)
                    self.image_label.config(image=self.photo)
                except Exception as e:
                    logging.error(f"无法显示图像: {e}")
                    self.image_label.config(image='')
            else:
                self.image_label.config(image='')

    def toggle_css_selector(self, *args):
        if self.enable_css_selector_var.get():
            self.css_selector_text.config(state=tk.NORMAL)
        else:
            self.css_selector_text.delete('1.0', tk.END)
            self.css_selector_text.config(state=tk.DISABLED)

    def toggle_tag_exclusion(self, *args):
        if self.enable_tag_exclusion_var.get():
            self.excluded_tags_text.config(state=tk.NORMAL)
        else:
            self.excluded_tags_text.delete('1.0', tk.END)
            self.excluded_tags_text.config(state=tk.DISABLED)

    def toggle_domain_exclusion(self, *args):
        if self.enable_domain_exclusion_var.get():
            self.excluded_domains_text.config(state=tk.NORMAL)
        else:
            self.excluded_domains_text.delete('1.0', tk.END)
            self.excluded_domains_text.config(state=tk.DISABLED)

    def toggle_custom_format(self):
        """切换自定义格式输入框的状态"""
        if self.enable_custom_format.get():
            self.custom_format_text.config(state=tk.NORMAL)
        else:
            self.custom_format_text.delete('1.0', tk.END)
            self.custom_format_text.config(state=tk.DISABLED)

    def _on_mousewheel(self, event, widget):
        """统一处理鼠标滚轮事件"""
        if isinstance(widget, (tk.Text, scrolledtext.ScrolledText)):
            if event.num == 5 or event.delta < 0:  # 向下滚动
                widget.yview_scroll(1, "units")
            elif event.num == 4 or event.delta > 0:  # 向上滚动
                widget.yview_scroll(-1, "units")
        elif isinstance(widget, tk.Listbox):
            if event.num == 5 or event.delta < 0:  # 向下滚动
                widget.yview_scroll(1, "units")
            elif event.num == 4 or event.delta > 0:  # 向上滚动
                widget.yview_scroll(-1, "units")
        return "break"  # 防止事件传播

    def process_html_content(self, content, options, is_cleaned=False):
        """处理HTML内容"""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            
            if is_cleaned:
                # 清理无用标签
                for tag in soup.find_all(['script', 'style', 'iframe', 'meta', 'link', 'noscript']):
                    tag.decompose()
                
                # 清理空标签
                for tag in soup.find_all():
                    if len(tag.get_text(strip=True)) == 0 and not tag.find(['img', 'video', 'audio']):
                        tag.decompose()
            
            # 根据选项处理元素
            if not options['preserve_images']:
                for img in soup.find_all('img'):
                    img.decompose()
            
            if not options['preserve_links']:
                for a in soup.find_all('a'):
                    a.replace_with(a.get_text(strip=True))
            
            if not options['preserve_tables']:
                for table in soup.find_all('table'):
                    # 将表格转换为文本
                    text = ' '.join(cell.get_text(strip=True) for cell in table.find_all(['td', 'th']))
                    table.replace_with(text)
            
            if not options['preserve_lists']:
                for list_tag in soup.find_all(['ul', 'ol']):
                    # 将列表转换为文本
                    text = ' '.join(li.get_text(strip=True) for li in list_tag.find_all('li'))
            if not options['preserve_code']:
                for code in soup.find_all(['code', 'pre']):
                    code.replace_with(code.get_text(strip=True))
            
            if not options['preserve_headings']:
                for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    # 将标题转换为普通段落
                    p = soup.new_tag('p')
                    p.string = heading.get_text(strip=True)
                    heading.replace_with(p)
            
            if not options['preserve_emphasis']:
                for em in soup.find_all(['em', 'strong', 'b', 'i']):
                    em.replace_with(em.get_text(strip=True))
            
            if not options['preserve_quotes']:
                for quote in soup.find_all(['blockquote', 'q']):
                    quote.replace_with(quote.get_text(strip=True))
            
            # 美化HTML输出
            if is_cleaned:
                # 添加基本样式
                style = soup.new_tag('style')
                style.string = """
                    body { font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }
                    img { max-width: 100%; height: auto; }
                    table { border-collapse: collapse; width: 100%; margin: 10px 0; }
                    td, th { border: 1px solid #ddd; padding: 8px; }
                    blockquote { border-left: 4px solid #ddd; margin: 0; padding-left: 20px; }
                    pre { background: #f5f5f5; padding: 15px; border-radius: 5px; overflow-x: auto; }
                    code { background: #f5f5f5; padding: 2px 5px; border-radius: 3px; }
                """
                soup.head.append(style)
            
            # 返回处理后的HTML
            return str(soup)
            
        except Exception as e:
            logging.error(f"HTML处理错误: {str(e)}")
            return content

    async def _save_json_async(self, data, file_path):
        """异步保存JSON数据"""
        try:
            async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
                await f.write(json.dumps(data, ensure_ascii=False, indent=2))
            self.saved_files.append(file_path)
            return file_path
        except Exception as e:
            logging.error(f"保存JSON数据失败: {e}")
            return None

    async def _save_screenshot_async(self, screenshot_data, url):
        """异步保存截图"""
        try:
            if not screenshot_data:
                return None
            
            filename = self.get_safe_filename(url)
            screenshot_path = self.directories['screenshots'] / f"{filename}.png"
            
            if isinstance(screenshot_data, str):
                # Base64 数据
                decoded_data = base64.b64decode(screenshot_data)
            elif isinstance(screenshot_data, bytes):
                decoded_data = screenshot_data
            else:
                return None
            
            async with aiofiles.open(screenshot_path, 'wb') as f:
                await f.write(decoded_data)
            
            self.saved_files.append(screenshot_path)
            return screenshot_path
        except Exception as e:
            logging.error(f"保存截图失败: {e}")
            return None

    def toggle_metadata_options(self):
        """切换元数据选项的状态"""
        state = 'normal' if self.extract_metadata_var.get() else 'disabled'
        for widget in self.metadata_options_frame.winfo_children():
            for child in widget.winfo_children():
                child.configure(state=state)

    def toggle_content_analysis(self):
        """切换内容分析选项的状态"""
        state = 'normal' if self.enable_content_analysis_var.get() else 'disabled'
        for widget in self.content_analysis_frame.winfo_children():
            widget.configure(state=state)

    def get_content_processing_config(self):
        """获取内容处理配置"""
        config = {
            'remove_noise': self.remove_noise_var.get(),
            'smart_extract': self.smart_extract_var.get(),
            'content_relevance_threshold': self.content_relevance_var.get(),
            'filter_options': {
                key: value.get() for key, value in self.filter_options.items()
            }
        }
        
        # 添加元数据配���
        if self.extract_metadata_var.get():
            config['metadata'] = {
                key.replace('extract_', ''): value.get()
                for key, value in self.metadata_options.items()
            }
        
        # 添加内容分析配置
        if self.enable_content_analysis_var.get():
            config['content_analysis'] = {
                key: value.get()
                for key, value in self.content_analysis_options.items()
            }
        
        return config

    async def save_browsable_page(self, html_content, url, resources=None):
        """保存完整的可浏览网页"""
        if not self.enable_page_clone.get():
            return None
            
        try:
            from bs4 import BeautifulSoup
            
            # 显示进度条和标签
            self.root.after(0, lambda: (
                self.progress_frame.pack(fill=tk.X, pady=5),
                self.progress_bar.pack(fill=tk.X),
                self.progress_label.pack(fill=tk.X)
            ))
            
            def update_progress(percentage, message):
                self.root.after(0, lambda: (
                    self.progress_var.set(percentage),
                    self.progress_label.config(text=message)
                ))
            
            update_progress(0, "准备克隆网页...")
            
            # 创建保存目录
            page_name = self.get_safe_filename(url)
            page_dir = self.base_dir / "pages" / page_name
            page_dir.mkdir(parents=True, exist_ok=True)
            
            update_progress(5, "解析HTML内容...")
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 创建资源目录结构
            resources_dir = page_dir / "resources"
            for subdir in ['css', 'js', 'images', 'fonts', 'media']:
                (resources_dir / subdir).mkdir(parents=True, exist_ok=True)
            
            # 收集所有需要下载的资源
            resource_items = []
            
            # CSS文件
            for css in soup.find_all('link', rel='stylesheet'):
                if css.get('href'):
                    resource_items.append(('css', css, 'href'))
            
            # JavaScript文件
            for js in soup.find_all('script', src=True):
                resource_items.append(('js', js, 'src'))
            
            # 图片文件
            for img in soup.find_all('img'):
                if img.get('src'):
                    resource_items.append(('images', img, 'src'))
                if img.get('srcset'):
                    srcset = img['srcset'].split(',')
                    for src in srcset:
                        url_part = src.strip().split()[0]
                        resource_items.append(('images', img, 'srcset', url_part))
            
            # 字体文件
            for font in soup.find_all('link', rel='font'):
                if font.get('href'):
                    resource_items.append(('fonts', font, 'href'))
            
            # 媒体文件
            for media in soup.find_all(['video', 'audio', 'source']):
                if media.get('src'):
                    resource_items.append(('media', media, 'src'))
            
            # 背景图片和其他CSS中的URL
            for style in soup.find_all(['style', 'link'], type='text/css'):
                if style.string:
                    urls = re.findall(r'url\([\'"]?([^\'"()]+)[\'"]?\)', style.string)
                    for url_match in urls:
                        resource_items.append(('images', style, 'style', url_match))
            
            # 内联样式中的URL
            for elem in soup.find_all(style=True):
                urls = re.findall(r'url\([\'"]?([^\'"()]+)[\'"]?\)', elem['style'])
                for url_match in urls:
                    resource_items.append(('images', elem, 'style', url_match))
            
            total_resources = len(resource_items)
            if total_resources == 0:
                total_resources = 1
            
            resources_processed = 0
            
            async def download_resource(resource_type, element, attr, url_part=None):
                """下载并保存资源文件"""
                try:
                    nonlocal resources_processed
                    
                    # 获取资源URL
                    resource_url = url_part if url_part else element.get(attr)
                    if not resource_url:
                        return None
                    
                    # 设置超和重试参数
                    timeout = aiohttp.ClientTimeout(total=30, connect=10)
                    max_retries = 3
                    retry_delay = 1  # 重试延迟（秒）

                    # 检查是否是 data URI
                    if resource_url.startswith('data:'):
                        try:
                            # 解析 data URI
                            header, data = resource_url.split(',', 1)
                            mime_type = header.split(';')[0].split(':')[1]
                            
                            # 生成文件名
                            file_name = f"data_uri_{hash(resource_url)}"
                            
                            # 确定文件扩展名
                            ext = mimetypes.guess_extension(mime_type) or f".{resource_type}"
                            
                            file_path = resources_dir / resource_type / f"{file_name}{ext}"
                            relative_path = file_path.relative_to(page_dir)
                            
                            # 检查是否已下载
                            if file_path.exists():
                                resources_processed += 1
                                progress = 10 + (resources_processed / total_resources * 80)
                                update_progress(progress, f"使用缓存: {file_name}")
                                return relative_path
                            
                            # 解码并保存数据
                            if ';base64,' in header:
                                decoded_data = base64.b64decode(data)
                            else:
                                # 处理URL编码的数据
                                from urllib.parse import unquote
                                decoded_data = unquote(data).encode('utf-8')
                            
                            async with aiofiles.open(file_path, 'wb') as f:
                                await f.write(decoded_data)
                            
                            resources_processed += 1
                            progress = 10 + (resources_processed / total_resources * 80)
                            update_progress(progress, f"保存Data URI: {file_name}")
                            
                            return relative_path
                            
                        except Exception as e:
                            logging.error(f"处理Data URI失败: {e}")
                            resources_processed += 1
                            return None
                    
                    # 处理常规URL
                    # 规范化URL
                    absolute_url = urljoin(url, resource_url)
                    parsed_url = urlparse(absolute_url)
                    
                    # 生成安全的文件名
                    file_name = re.sub(r'[<>:"/\\|?*]', '_', parsed_url.path.split('/')[-1])
                    if not file_name:
                        file_name = f"resource_{hash(absolute_url)}"
                    
                    # 确定文件扩展名
                    ext = mimetypes.guess_extension(mimetypes.guess_type(file_name)[0] or '')
                    if not ext:
                        ext = f".{resource_type}"
                    
                    file_path = resources_dir / resource_type / f"{file_name}{ext}"
                    relative_path = file_path.relative_to(page_dir)
                    
                    # 检查是否已下载
                    if file_path.exists():
                        resources_processed += 1
                        progress = 10 + (resources_processed / total_resources * 80)
                        update_progress(progress, f"使用缓存: {file_name}")
                        return relative_path
                    
                    # 如果资源已存在于resources字典中
                    if resources and absolute_url in resources:
                        async with aiofiles.open(file_path, 'wb') as f:
                            await f.write(resources[absolute_url])
                    else:
                        # 添加重试逻辑
                        for retry in range(max_retries):
                            try:
                                async with aiohttp.ClientSession(timeout=timeout) as session:
                                    async with session.get(absolute_url, ssl=False) as response:  # 禁用SSL验证
                                        if response.status == 200:
                                            content = await response.read()
                                            async with aiofiles.open(file_path, 'wb') as f:
                                                await f.write(content)
                                            break  # 下载成功，跳出重试循环
                                        elif response.status == 404:
                                            logging.warning(f"资源不存在: {absolute_url}")
                                            break  # 资源不存在，不需要重试
                                        else:
                                            raise aiohttp.ClientError(f"HTTP {response.status}")
                            except Exception as e:
                                if retry < max_retries - 1:
                                    logging.warning(f"下载失败，正在重试 ({retry + 1}/{max_retries}): {absolute_url}")
                                    await asyncio.sleep(retry_delay * (retry + 1))  # 递增延迟
                                else:
                                    logging.error(f"下载失败 {absolute_url}: {str(e)}")
                                    return None
                    
                    resources_processed += 1
                    progress = 10 + (resources_processed / total_resources * 80)
                    update_progress(progress, f"下载资源: {file_name}")
                    
                    return relative_path
                    
                except Exception as e:
                    logging.error(f"下载资源失败 {resource_url}: {e}")
                    resources_processed += 1
                    return None
            
            # 并发下载所有资源
            semaphore = asyncio.Semaphore(5)  # 限制并发下载数量
            
            async def controlled_download(*args):
                async with semaphore:
                    return await download_resource(*args)

            # 修改任务创建部分
            tasks = []
            for item in resource_items:
                if len(item) == 4:
                    resource_type, element, attr, url_part = item
                    tasks.append((item, controlled_download(resource_type, element, attr, url_part)))
                else:
                    resource_type, element, attr = item
                    tasks.append((item, controlled_download(resource_type, element, attr)))
            
            # 等待所有下载任务完成
            update_progress(10, "开始下载资源...")
            results = []
            for item, task in tasks:
                path = await task
                if path:
                    results.append((item, path))
            
            # 更新HTML中的资源路径
            update_progress(90, "更新资源路径...")
            for item, path in results:
                if len(item) == 4:  # srcset或style中的URL
                    resource_type, element, attr, url_part = item
                    if attr == 'style':
                        element['style'] = element['style'].replace(url_part, str(path))
                    elif attr == 'srcset':
                        srcset = element['srcset'].split(',')
                        new_srcset = []
                        for src in srcset:
                            if url_part in src:
                                new_srcset.append(src.replace(url_part, str(path)))
                            else:
                                new_srcset.append(src)
                        element['srcset'] = ', '.join(new_srcset)
                else:
                    resource_type, element, attr = item
                    element[attr] = str(path)
            
            # 添加基础样式和元数据
            update_progress(95, "添加元数据...")
            if not soup.head:
                soup.html.insert(0, soup.new_tag('head'))
            
            # 添加元数据
            meta_tags = {
                'charset': 'utf-8',
                'viewport': 'width=device-width, initial-scale=1',
                'description': '克隆的网页',
                'generator': 'Tai-网页爬虫'
            }
            
            for name, content in meta_tags.items():
                meta = soup.new_tag('meta')
                if name == 'charset':
                    meta['charset'] = content
                else:
                    meta['name'] = name
                    meta['content'] = content
                soup.head.insert(0, meta)
            
            # 保存完整的HTML
            update_progress(98, "保存HTML文件...")
            html_path = page_dir / "index.html"
            async with aiofiles.open(html_path, 'w', encoding='utf-8') as f:
                await f.write(str(soup.prettify()))
            
            update_progress(100, "网页克隆完成!")
            
            # 3秒后隐藏进度条
            self.root.after(3000, lambda: self.progress_frame.pack_forget())
            
            self.saved_files.append(html_path)
            return html_path
            
        except Exception as e:
            logging.error(f"保存可浏览网页失败: {e}")
            update_progress(100, f"克隆失败: {str(e)}")
            self.root.after(3000, lambda: self.progress_frame.pack_forget())
            return None

    async def extract_pure_text(self, html_content):
        """提取网页纯文本内容"""
        if not self.enable_text_extract.get():
            return None
        
        try:
            from bs4 import BeautifulSoup
            import re
            from collections import OrderedDict
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            from PIL import Image

            # 显示进度条
            self.root.after(0, lambda: (
                self.progress_frame.pack(fill=tk.X, pady=5),
                self.progress_bar.pack(fill=tk.X),
                self.progress_label.pack(fill=tk.X)
            ))

            def update_progress(percentage, message):
                self.root.after(0, lambda: (
                    self.progress_var.set(percentage),
                    self.progress_label.config(text=message)
                ))

            update_progress(0, "准备提取纯文本...")

            # 解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            update_progress(20, "移除无用元素...")
            
            # 移除脚本和样式
            for element in soup(['script', 'style', 'noscript']):
                element.decompose()

            # 根据选项移除元素
            if self.text_extract_options['remove_ads'].get():
                for element in soup.find_all(class_=re.compile(r'ad|banner|sponsor|commercial', re.I)):
                    element.decompose()

            if self.text_extract_options['remove_menus'].get():
                for element in soup.find_all(['nav', 'menu']):
                    element.decompose()
                for element in soup.find_all(class_=re.compile(r'menu|nav|navigation', re.I)):
                    element.decompose()

            if self.text_extract_options['remove_headers'].get():
                for element in soup.find_all(['header']):
                    element.decompose()
                for element in soup.find_all(class_=re.compile(r'header|top-bar', re.I)):
                    element.decompose()

            if self.text_extract_options['remove_footers'].get():
                for element in soup.find_all(['footer']):
                    element.decompose()
                for element in soup.find_all(class_=re.compile(r'footer|bottom', re.I)):
                    element.decompose()

            # 识别主要内容区域
            if self.text_extract_options['keep_main_content'].get():
                main_content = None
                # 尝试找到主要内容区域
                for selector in ['main', 'article', '#content', '.content', '#main', '.main']:
                    main_content = soup.select_one(selector)
                    if main_content:
                        soup = BeautifulSoup(str(main_content), 'html.parser')
                        break

            update_progress(40, "提取内容...")

            # 创建Word文档
            doc = Document()
            
            # 设置文档标题
            if hasattr(soup.find('title'), 'text'):
                title = soup.find('title').text.strip()
                heading = doc.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 处理内容
            elements = []
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'img', 'table', 'a', 'ul', 'ol']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(element.name[1])
                    text = element.get_text(strip=True)
                    if text:
                        elements.append(('heading', level, text))
                
                elif element.name == 'p':
                    text = element.get_text(strip=True)
                    if text and len(text) > 20:
                        elements.append(('paragraph', text))
                
                elif element.name == 'img' and self.text_extract_options['keep_images'].get():
                    src = element.get('src', '')
                    alt = element.get('alt', '图片')
                    if src:
                        elements.append(('image', src, alt))
                
                elif element.name == 'table' and self.text_extract_options['keep_tables'].get():
                    elements.append(('table', element))
                
                elif element.name == 'a' and self.text_extract_options['keep_links'].get():
                    text = element.get_text(strip=True)
                    href = element.get('href', '')
                    if text and href:
                        elements.append(('link', text, href))
                
                elif element.name in ['ul', 'ol']:
                    items = []
                    for li in element.find_all('li'):
                        text = li.get_text(strip=True)
                        if text:
                            items.append(text)
                    if items:
                        elements.append(('list', element.name, items))

            update_progress(60, "处理格式...")

            # 处理元素
            for element in elements:
                if element[0] == 'heading':
                    heading = doc.add_heading('', element[1])
                    heading.add_run(element[2])
                
                elif element[0] == 'paragraph':
                    para = doc.add_paragraph()
                    para.add_run(element[1])
                
                elif element[0] == 'image':
                    try:
                        response = await self.download_image(element[1])
                        if response:
                            doc.add_picture(io.BytesIO(response), width=Pt(300))
                            if element[2]:
                                doc.add_paragraph(element[2], style='Caption')
                    except Exception as e:
                        logging.error(f"图片处理失败: {e}")
                
                elif element[0] == 'table':
                    rows = element[1].find_all('tr')
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                        for i, row in enumerate(rows):
                            for j, cell in enumerate(row.find_all(['td', 'th'])):  # 添加缺少的右括号
                                table.cell(i, j).text = cell.get_text(strip=True)
                
                elif element[0] == 'link':
                    para = doc.add_paragraph()
                    run = para.add_run(f"{element[1]} ({element[2]})")
                    run.font.color.rgb = RGBColor(0, 0, 255)
                
                elif element[0] == 'list':
                    for item in element[2]:
                        para = doc.add_paragraph()
                        para.style = 'List Bullet' if element[1] == 'ul' else 'List Number'
                        para.add_run(item)

            update_progress(80, "保存文件...")

            # 保存文件
            file_name = self.get_safe_filename(self.url_var.get())
            if self.text_extract_options['save_as_word'].get():
                file_path = self.base_dir / "text" / f"{file_name}_content.docx"
                doc.save(str(file_path))
            else:
                # 保存为纯文本
                file_path = self.base_dir / "text" / f"{file_name}_content.txt"
                text_content = []
                for element in elements:
                    if element[0] == 'heading':
                        text_content.append(f"\n{'#' * element[1]} {element[2]}\n")
                    elif element[0] == 'paragraph':
                        text_content.append(element[1])
                    elif element[0] == 'link':
                        text_content.append(f"{element[1]} <{element[2]}>")
                    elif element[0] == 'list':
                        for i, item in enumerate(element[2], 1):
                            text_content.append(f"{'*' if element[1] == 'ul' else str(i)+'.'} {item}")
                
                async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
                    await f.write('\n\n'.join(text_content))

            update_progress(100, "提取完成!")
            self.root.after(3000, lambda: self.progress_frame.pack_forget())
            
            self.saved_files.append(file_path)
            return file_path

        except ImportError as e:
            logging.error(f"导入所需模块失败: {e}")
            return None
        except Exception as e:
            logging.error(f"提取纯文本失败: {e}")
            update_progress(100, f"提取失败: {str(e)}")
            self.root.after(3000, lambda: self.progress_frame.pack_forget())
            return None

    async def download_image(self, url):
        """下载图片"""
        try:
            if url.startswith('data:'):
                # 处理 base64 图片
                header, data = url.split(',', 1)
                return base64.b64decode(data)
            else:
                # 下载网络图片
                async with aiohttp.ClientSession() as session:
                    async with session.get(url, ssl=False) as response:
                        if response.status == 200:
                            return await response.read()
        except Exception as e:
            logging.error(f"下载图片失败 {url}: {e}")
            return None

    def toggle_llm_optimize(self):
        """切换LLM优化选项的状态"""
        if self.enable_text_extract.get():
            self.llm_optimize_cb.configure(state=tk.NORMAL)
        else:
            self.enable_llm_optimize.set(False)
            self.llm_optimize_cb.configure(state=tk.DISABLED)
            self.llm_settings_frame.pack_forget()

    def toggle_llm_settings(self):
        """切换LLM设置框架的��示状态"""
        if self.enable_llm_optimize.get():
            self.llm_settings_frame.pack(fill=tk.X, pady=5)
            # 根据当前选择的模型类型显示对应设置
            self.toggle_model_settings()
        else:
            self.llm_settings_frame.pack_forget()

    def toggle_model_settings(self):
        """切换模型设置框架的显示"""
        if not self.enable_llm_optimize.get():
            return
        
        if self.model_type_var.get() == "local":
            self.local_model_frame.pack(fill=tk.X, padx=5, pady=5)
            self.api_model_frame.pack_forget()
        else:
            self.local_model_frame.pack_forget()
            self.api_model_frame.pack(fill=tk.X, padx=5, pady=5)

    async def optimize_with_llm(self, text, model_name):
        """使用LLM优化文本"""
        try:
            logging.info(f"开始LLM优化，使用模型: {model_name}")
            logging.info("当前启用的优化选项:")
            for option, var in self.llm_optimize_options.items():
                logging.info(f"  - {option}: {var.get()}")
            
            # 检查模型目录
            models_dir = Path("models")
            if not models_dir.exists():
                logging.info("创建模型目录")
                models_dir.mkdir(parents=True, exist_ok=True)
            
            # 构建提示词
            if self.enable_custom_prompt.get():
                custom_prompt = self.custom_prompt_text.get('1.0', tk.END).strip()
                if custom_prompt:
                    prompt = custom_prompt + "\n\n" + text
                else:
                    # 使用默认提示词
                    prompts = []
                    if self.llm_optimize_options['improve_readability'].get():
                        prompts.append("提高文本的可读性和流畅度")
                    if self.llm_optimize_options['enhance_structure'].get():
                        prompts.append("优化文本结构和段落组织")
                    if self.llm_optimize_options['fix_grammar'].get():
                        prompts.append("修正语法错误和表达不准确的地方")
                    if self.llm_optimize_options['summarize'].get():
                        prompts.append("为文本生成简短的摘要")
                    if self.llm_optimize_options['translate'].get():
                        prompts.append("将文本翻译成英文")
                    prompt = "请" + "、".join(prompts) + "。以下是原文：\n\n" + text
            else:
                # 使用默认提示词
                prompts = []
                if self.llm_optimize_options['improve_readability'].get():
                    prompts.append("提高文本的可读性和流畅度")
                if self.llm_optimize_options['enhance_structure'].get():
                    prompts.append("优化文本结构和段落组织")
                if self.llm_optimize_options['fix_grammar'].get():
                    prompts.append("修正语法错误和表达不准确的地方")
                if self.llm_optimize_options['summarize'].get():
                    prompts.append("为文本生成简短的摘要")
                if self.llm_optimize_options['translate'].get():
                    prompts.append("将文本翻译成英文")
                prompt = "请" + "、".join(prompts) + "。以下是原文：\n\n" + text

            logging.debug(f"生成的提示词: {prompt[:200]}...")

            try:
                from llama_cpp import Llama
                logging.info("成功导入llama-cpp模块")
            except ImportError as e:
                logging.error("导入llama-cpp模块失败，尝试安装...")
                import subprocess
                try:
                    subprocess.check_call([sys.executable, "-m", "pip", "install", "llama-cpp-python"])
                    from llama_cpp import Llama
                    logging.info("成功安装并导入llama-cpp模块")
                except Exception as install_error:
                    logging.error(f"安装llama-cpp-python失败: {install_error}")
                    raise

            # 修改模型路径映射和配置逻辑
            model_paths = {
                "chatglm3-6b": models_dir / "chatglm3-6b.Q4_K_M.gguf",
                "llama-2-7b": models_dir / "llama-2-7b.Q4_K_M.gguf",
                "baichuan2-7b": models_dir / "baichuan2-7b.Q4_K_M.gguf",
                "qwen-7b": models_dir / "qwen-7b.Q4_K_M.gguf",
                "yi-6b": models_dir / "yi-6b.Q4_K_M.gguf",
                "deepseek-7b": models_dir / "deepseek-7b.Q4_K_M.gguf"
            }

            # 规范化模型名称
            normalized_model_name = model_name.lower().replace('_', '-').split('.')[0]
            
            # 获取模型路径
            model_path = None
            for key, path in model_paths.items():
                if key in normalized_model_name:
                    model_path = path
                    break

            if not model_path:
                error_msg = f"未找到模型配置: {model_name}\n支持的模型: {', '.join(model_paths.keys())}"
                logging.error(error_msg)
                # 显示错误对话框
                self.root.after(0, lambda: messagebox.showerror(
                    "错误",
                    f"未找到模型配置。\n请在模型管理器中下载以下模型之一：\n{', '.join(model_paths.keys())}"
                ))
                raise ValueError(error_msg)
            
            # 检查模型文件是否存在
            if not model_path.exists():
                error_msg = f"模型文件不存在: {model_path}"
                logging.error(error_msg)
                # 显示错误对话框
                self.root.after(0, lambda: messagebox.showerror(
                    "错误",
                    f"模型文件不存在。\n请在模型管理器中下载 {model_name} 模型。"
                ))
                raise FileNotFoundError(error_msg)
            
            logging.info(f"使用模型文件: {model_path}")
            logging.info(f"模型文件大小: {model_path.stat().st_size / 1024 / 1024:.2f} MB")
            
            # 初始化模型
            try:
                llm = Llama(model_path=str(model_path), n_ctx=4096)
                logging.info("模型加载完成")
            except Exception as e:
                error_msg = f"模型加载失败: {e}"
                logging.error(error_msg)
                self.root.after(0, lambda: messagebox.showerror("错误", error_msg))
                raise
            
            # 获取响应
            logging.info("开始生成响应...")
            try:
                response = llm(prompt, max_tokens=4096, temperature=0.7)
                logging.info("响应生成完成")
            except Exception as e:
                error_msg = f"响应生成失败: {e}"
                logging.error(error_msg)
                self.root.after(0, lambda: messagebox.showerror("错误", error_msg))
                raise
            
            result = response['choices'][0]['text']
            logging.debug(f"���成的响应: {result[:200]}...")
            
            return result

        except Exception as e:
            logging.error(f"LLM处理失败: {e}", exc_info=True)
            # 显示错误对话框
            self.root.after(0, lambda: messagebox.showerror(
                "LLM处理失败",
                f"处理过程中发生错误：\n{str(e)}\n\n请确保已正确安装并下载所需模型。"
            ))
            raise

    def start_auto_download(self, dialog, model_name, model_path):
        """开始自动下载"""
        dialog.destroy()  # 关闭提示对话框
        asyncio.create_task(self._download_model(model_name, model_path))  # 启动下载

    async def _download_model(self, model_name, model_path):
        """下载模型文件"""
        try:
            logging.info(f"开始下载模型: {model_name}")
            logging.debug(f"目标路径: {model_path}")
            
            # 检查目标目录是否存在
            if not model_path.parent.exists():
                logging.info(f"创建目录: {model_path.parent}")
                model_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 创建下载进度窗口
            progress_window = tk.Toplevel(self.root)
            progress_window.title(f"下载模型 - {model_name}")
            progress_window.geometry("400x200")
            progress_window.transient(self.root)
            progress_window.grab_set()  # 使窗口模态
            progress_window.focus_set()  # 获取焦点
            
            # 将窗口居中
            window_width = 400
            window_height = 200
            screen_width = self.root.winfo_screenwidth()
            screen_height = self.root.winfo_screenheight()
            x = (screen_width - window_width) // 2
            y = (screen_height - window_height) // 2
            progress_window.geometry(f"{window_width}x{window_height}+{x}+{y}")
            
            # 添加进度信息标签
            info_label = ttk.Label(progress_window, text="准备下载...", wraplength=380)
            info_label.pack(pady=10)
            
            # 添加文件大小信息标签
            size_label = ttk.Label(progress_window, text="")
            size_label.pack(pady=5)
            
            # 添加下载速度标签
            speed_label = ttk.Label(progress_window, text="")
            speed_label.pack(pady=5)
            
            # 添加进度条
            progress_var = tk.DoubleVar(value=0.0)
            progress_bar = ttk.Progressbar(
                progress_window,
                variable=progress_var,
                maximum=100,
                mode='determinate',
                length=350
            )
            progress_bar.pack(pady=10)
            
            # 添加取消按钮
            cancel_var = tk.BooleanVar(value=False)
            cancel_button = ttk.Button(
                progress_window,
                text="取消下载",
                command=lambda: cancel_var.set(True)
            )
            cancel_button.pack(pady=10)
            
            # 获取模型URL
            # 定义可用模型列表
            available_models = {
                "ChatGLM3-6B": {
                    "name": "chatglm3-6b.Q4_K_M.gguf",
                    "url": "https://huggingface.co/TheBloke/ChatGLM3-6B-GGUF/resolve/main/chatglm3-6b.Q4_K_M.gguf"
                },
                "Llama-2-7B": {
                    "name": "llama-2-7b.Q4_K_M.gguf",
                    "url": "https://huggingface.co/TheBloke/Llama-2-7B-GGUF/resolve/main/llama-2-7b.Q4_K_M.gguf"
                },
                "Qwen-7B": {
                    "name": "qwen-7b.Q4_K_M.gguf",
                    "url": "https://huggingface.co/TheBloke/Qwen-7B-GGUF/resolve/main/qwen-7b.Q4_K_M.gguf"
                },
                "Yi-6B": {
                    "name": "yi-6b.Q4_K_M.gguf",
                    "url": "https://huggingface.co/TheBloke/Yi-6B-GGUF/resolve/main/yi-6b.Q4_K_M.gguf"
                },
                "Mistral-7B": {
                    "name": "mistral-7b.Q4_K_M.gguf",
                    "url": "https://huggingface.co/TheBloke/Mistral-7B-GGUF/resolve/main/mistral-7b.Q4_K_M.gguf"
                },
                "Neural-7B": {
                    "name": "neural-7b.Q4_K_M.gguf",
                    "url": "https://huggingface.co/TheBloke/Neural-7B-GGUF/resolve/main/neural-7b.Q4_K_M.gguf"
                }
            }

            # 查找模型URL
            model_url = None
            for model_info in available_models.values():
                if model_info['name'].startswith(model_name):
                    model_url = model_info['url']
                    break

            if not model_url:
                error_msg = f"未找到模型 {model_name} 的下载链接"
                logging.error(error_msg)
                raise ValueError(error_msg)
            
            logging.info(f"下载URL: {model_url}")
            info_label.config(text=f"正在连接到服务器...\n{model_url}")
            
            # 创建临时文件
            temp_path = model_path.with_suffix('.tmp')
            logging.debug(f"临时文件路径: {temp_path}")
            
            try:
                logging.info("创建下载会话...")
                async with aiohttp.ClientSession() as session:
                    logging.info("开始下载请求...")
                    async with session.get(model_url) as response:
                        if response.status != 200:
                            error_msg = f"下载失败: HTTP {response.status}"
                            logging.error(error_msg)
                            raise aiohttp.ClientError(error_msg)
                        
                        total_size = int(response.headers.get('content-length', 0))
                        total_size_mb = total_size / (1024 * 1024)
                        logging.info(f"文件大小: {total_size_mb:.2f} MB")
                        
                        size_label.config(text=f"文件大小: {total_size_mb:.2f} MB")
                        info_label.config(text=f"正在下载模型: {model_name}")
                        
                        # 下载并显示进度
                        chunk_size = 1024 * 1024  # 1MB
                        downloaded = 0
                        start_time = datetime.now()
                        
                        async with aiofiles.open(temp_path, 'wb') as f:
                            async for chunk in response.content.iter_chunked(chunk_size):
                                if cancel_var.get():
                                    raise Exception("用户取消下载")
                                
                                await f.write(chunk)
                                downloaded += len(chunk)
                                
                                # 计算进度
                                progress = (downloaded / total_size) * 100
                                downloaded_mb = downloaded / (1024 * 1024)
                                
                                # 计算下载速度
                                elapsed_time = (datetime.now() - start_time).total_seconds()
                                if elapsed_time > 0:
                                    speed = downloaded / (1024 * 1024 * elapsed_time)  # MB/s
                                    eta = (total_size - downloaded) / (downloaded / elapsed_time)
                                    eta_str = str(timedelta(seconds=int(eta)))
                                    
                                    # 更新界面
                                    progress_var.set(progress)
                                    speed_label.config(text=f"下载速度: {speed:.2f} MB/s\n预计剩余时间: {eta_str}")
                                    size_label.config(text=f"已下载: {downloaded_mb:.2f} MB / {total_size_mb:.2f} MB")
                                
                                # 更新日志
                                if downloaded % (50 * chunk_size) == 0:  # 每50MB记录一次日志
                                    logging.info(f"下载进度: {progress:.1f}% ({downloaded_mb:.1f} MB / {total_size_mb:.1f} MB)")
                                
                                # 处理GUI事件
                                progress_window.update()
                
                # 下载���成后重命名文件
                logging.info("下载完成，重命名临时文件...")
                temp_path.rename(model_path)
                logging.info(f"模型文件已保存到: {model_path}")
                
                # 关闭进度窗口
                self.root.after(1000, progress_window.destroy)
                
            except Exception as e:
                logging.error(f"下载失败: {e}")
                if temp_path.exists():
                    logging.info("删除临时文件...")
                    temp_path.unlink()
                
                # 显示错误信息
                info_label.config(text=f"下载失败: {str(e)}")
                progress_window.update()
                
                # 3秒后关闭窗口
                self.root.after(3000, progress_window.destroy)
                raise
                
        except Exception as e:
            logging.error(f"下载过程发生错误: {e}")
            # 显示错误对话框
            self.root.after(0, lambda: messagebox.showerror(
                "下载失败",
                f"下载过程中发生错误：\n{str(e)}"
            ))
            raise

    def batch_optimize_files(self):
        """启动批量优化文件的异步操作"""
        def run_async():
            asyncio.run(self._batch_optimize_files())
        
        # 在新线程中运行异步操作
        threading.Thread(target=run_async).start()

    async def _batch_optimize_files(self):
        """批量优化文件的异步实现"""
        try:
            logging.info("开始批量文件优化")
            from tkinter import filedialog
            
            # 在主线程中选择文件
            file_paths = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: filedialog.askopenfilenames(
                    title="选择要优化的文件",
                    filetypes=[("Word文档", "*.docx"), ("文本文件", "*.txt")],
                    initialdir=self.base_dir / "text"
                )
            )
            
            if not file_paths:
                logging.info("用户取消了文件选择")
                return

            logging.info(f"选择的文件数量: {len(file_paths)}")
            for path in file_paths:
                logging.info(f"文件: {path}")

            # 显示进度条
            self.progress_frame.pack(fill=tk.X, pady=5)
            self.progress_bar.pack(fill=tk.X)
            self.progress_label.pack(fill=tk.X)
            
            def update_progress(percentage, message):
                self.root.after(0, lambda: (
                    self.progress_var.set(percentage),
                    self.progress_label.config(text=message)
                ))
                logging.info(f"进度: {percentage}% - {message}")

            # 读取所有文件内容
            all_content = []
            total_size = 0
            for i, file_path in enumerate(file_paths):
                update_progress(i / len(file_paths) * 40, f"读取文件: {Path(file_path).name}")
                
                file_size = Path(file_path).stat().st_size
                total_size += file_size
                logging.info(f"处理文件 {file_path}, 大小: {file_size/1024:.2f} KB")
                
                try:
                    if file_path.endswith('.docx'):
                        from docx import Document
                        doc = Document(file_path)
                        content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                        logging.info(f"成功读取Word文档，段落数: {len(doc.paragraphs)}")
                    else:
                        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                            content = await f.read()
                        logging.info(f"成功读取文本文件，字符数: {len(content)}")
                    
                    all_content.append(content)
                except Exception as e:
                    logging.error(f"读取文件失败 {file_path}: {e}", exc_info=True)
                    raise

            logging.info(f"总文件大小: {total_size/1024:.2f} KB")
            
            # 合并所有内容
            combined_content = "\n\n=== 分隔线 ===\n\n".join(all_content)
            logging.info(f"合并后的内容大小: {len(combined_content)} 字符")
            
            # 使用LLM优化
            update_progress(50, "正在使用LLM优化文本...")
            try:
                optimized_content = await self.optimize_with_llm(combined_content, self.llm_model_var.get())
                logging.info(f"LLM优化完成，优化后内容大小: {len(optimized_content)} 字符")
            except Exception as e:
                logging.error("LLM优化失败", exc_info=True)
                raise
            
            # 保存优化后的内容
            update_progress(80, "保存优化的文本...")
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.base_dir / "text" / f"optimized_{timestamp}.docx"
            
            try:
                # 创建新的Word文档
                doc = Document()
                doc.add_heading('LLM优化后的文档', 0)
                
                # 添加优化信息
                doc.add_paragraph(f"优化时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"使用模型: {self.llm_model_var.get()}")
                
                enabled_options = [name for name, var in self.llm_optimize_options.items() if var.get()]
                doc.add_paragraph("优化选项: " + ", ".join(enabled_options))
                
                # 添加源文件信息
                doc.add_paragraph("源文件:")
                for path in file_paths:
                    doc.add_paragraph(f"  - {Path(path).name}")
                
                # 添加优化后的内容
                doc.add_paragraph("\n" + optimized_content)
                
                # 保存文档
                doc.save(str(output_path))
                logging.info(f"文档已保存: {output_path}")
                
            except Exception as e:
                logging.error("保存文档失败", exc_info=True)
                raise
            
            update_progress(100, "优化完成!")
            self.root.after(3000, lambda: self.progress_frame.pack_forget())
            
            # 显示完成消息
            self.content_text.insert(tk.END, f"\n\nLLM优化完成，文件已保存至: {output_path}")
            
            # 添加到文件列表
            self.saved_files.append(output_path)
            self.files_listbox.insert(tk.END, str(output_path))
            
            logging.info("批量优化过程完成")

        except Exception as e:
            logging.error("批量优化失败", exc_info=True)
            self.content_text.insert(tk.END, f"\n\n批量优化失败: {str(e)}")
            update_progress(100, f"优化失败: {str(e)}")
            self.root.after(3000, lambda: self.progress_frame.pack_forget())

    def setup_logging(self):
        """配置日志系统"""
        # 创建logs目录
        log_dir = Path("logs")
        log_dir.mkdir(exist_ok=True)
        
        # 创建带时间戳的日志文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        log_file = log_dir / f"crawler_{timestamp}.log"
        
        # 配置日志格式
        formatter = logging.Formatter(
            '%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        
        # 文件处理器
        file_handler = logging.FileHandler(log_file, encoding='utf-8')
        file_handler.setFormatter(formatter)
        file_handler.setLevel(logging.DEBUG)
        
        # 控制台处理器
        console_handler = logging.StreamHandler()
        console_handler.setFormatter(formatter)
        console_handler.setLevel(logging.INFO)
        
        # 配置根日志记录器
        logger = logging.getLogger()
        logger.setLevel(logging.DEBUG)
        logger.addHandler(file_handler)
        logger.addHandler(console_handler)
        
        logging.info("日志系统初始化完成")

    def scan_local_models(self):
        """扫描本地已下载的模型文件"""
        models_dir = Path("models")
        if not models_dir.exists():
            models_dir.mkdir(parents=True, exist_ok=True)
            return []
        
        # 查找所有.gguf文件
        model_files = list(models_dir.glob("*.gguf"))
        return [model.name for model in model_files]

    def show_model_manager(self):
        """显示模型管理器"""
        manager_window = tk.Toplevel(self.root)
        manager_window.title("模型管理器")
        manager_window.geometry("800x600")
        
        # 使窗口模态
        manager_window.transient(self.root)
        manager_window.grab_set()
        
        # 创建主框架
        main_frame = ttk.Frame(manager_window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建模型列表框���
        list_frame = ttk.LabelFrame(main_frame, text="已安装的模型", padding="5")
        list_frame.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        
        # 创建模型列表
        model_list = tk.Listbox(list_frame, width=40)
        model_list.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        
        # 添加滚动条
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=model_list.yview)
        scrollbar.pack(fill=tk.Y, side=tk.RIGHT)
        model_list.configure(yscrollcommand=scrollbar.set)
        
        # 刷新本地模型列��
        def refresh_model_list():
            model_list.delete(0, tk.END)
            local_models = self.scan_local_models()
            for model in local_models:
                model_list.insert(tk.END, model)
        
        # 创建可用模型列表
        available_models = {
            "ChatGLM3-6B": {
                "name": "chatglm3-6b.Q4_K_M.gguf",
                "size": "3.1GB",
                "description": "智谱AI开源的对话模型，支持中英双语",
                "url": "https://huggingface.co/TheBloke/ChatGLM3-6B-GGUF/resolve/main/chatglm3-6b.Q4_K_M.gguf"
            },
            "Llama-2-7B": {
                "name": "llama-2-7b.Q4_K_M.gguf",
                "size": "3.8GB",
                "description": "Meta开源的大语言模型，性能优秀",
                "url": "https://huggingface.co/TheBloke/Llama-2-7B-GGUF/resolve/main/llama-2-7b.Q4_K_M.gguf"
            },
            "Qwen-7B": {
                "name": "qwen-7b.Q4_K_M.gguf",
                "size": "3.8GB",
                "description": "阿里云开源的通用大语言模型",
                "url": "https://huggingface.co/TheBloke/Qwen-7B-GGUF/resolve/main/qwen-7b.Q4_K_M.gguf"
            },
            "Yi-6B": {
                "name": "yi-6b.Q4_K_M.gguf",
                "size": "3.1GB",
                "description": "零一万物开源的大语言模型",
                "url": "https://huggingface.co/TheBloke/Yi-6B-GGUF/resolve/main/yi-6b.Q4_K_M.gguf"
            },
            "Mistral-7B": {
                "name": "mistral-7b.Q4_K_M.gguf",
                "size": "3.8GB",
                "description": "Mistral AI开源的高性能模型",
                "url": "https://huggingface.co/TheBloke/Mistral-7B-GGUF/resolve/main/mistral-7b.Q4_K_M.gguf"
            },
            "Neural-7B": {
                "name": "neural-7b.Q4_K_M.gguf",
                "size": "3.8GB",
                "description": "Neural开源的高性能中文模型",
                "url": "https://huggingface.co/TheBloke/Neural-7B-GGUF/resolve/main/neural-7b.Q4_K_M.gguf"
            }
        }
        
        # 创建右侧信息框架
        info_frame = ttk.Frame(main_frame)
        info_frame.pack(fill=tk.BOTH, expand=True, side=tk.RIGHT, padx=10)
        
        # 创建可用模型下拉框
        model_var = tk.StringVar()
        ttk.Label(info_frame, text="选择要下载的模型:").pack(anchor=tk.W, pady=(0, 5))
        model_combo = ttk.Combobox(info_frame, textvariable=model_var, values=list(available_models.keys()))
        model_combo.pack(fill=tk.X, pady=(0, 10))
        
        # 创建模型信息显示区域
        info_text = scrolledtext.ScrolledText(info_frame, height=10, width=50)
        info_text.pack(fill=tk.BOTH, expand=True)
        
        def update_model_info(*args):
            info_text.delete(1.0, tk.END)
            model = available_models.get(model_var.get())
            if model:
                info_text.insert(tk.END, f"模型名称: {model['name']}\n")
                info_text.insert(tk.END, f"文件大小: {model['size']}\n")
                info_text.insert(tk.END, f"模型描述: {model['description']}\n")
                info_text.insert(tk.END, f"\n下载地址: {model['url']}\n")
        
        model_var.trace('w', update_model_info)
        
        # 创建按钮框架
        button_frame = ttk.Frame(info_frame)
        button_frame.pack(fill=tk.X, pady=10)
        
        # 复制链接按钮
        def copy_url():
            model = available_models.get(model_var.get())
            if model:
                manager_window.clipboard_clear()
                manager_window.clipboard_append(model['url'])
                messagebox.showinfo("提示", "链接已复制到剪贴板")
        
        ttk.Button(button_frame, text="复制链接", 
                  command=copy_url).pack(side=tk.LEFT, padx=5)
        
        # 浏览器打开按钮
        def open_in_browser():
            model = available_models.get(model_var.get())
            if model:
                webbrowser.open(model['url'])
        
        ttk.Button(button_frame, text="在浏览器中打开", 
                  command=open_in_browser).pack(side=tk.LEFT, padx=5)
        
        # 自动下载按钮
        def start_download():
            model = available_models.get(model_var.get())
            if model:
                manager_window.destroy()
                self.start_async_download(
                    model['name'].split('.')[0],
                    Path("models") / model['name']
                )
        
        ttk.Button(button_frame, text="自动下载", 
                  command=start_download).pack(side=tk.LEFT, padx=5)
        
        # 删除按钮
        def delete_model():
            selection = model_list.curselection()
            if selection:
                model_name = model_list.get(selection[0])
                if messagebox.askyesno("确认", f"确定要删除模型 {model_name} 吗？"):
                    try:
                        model_path = Path("models") / model_name
                        if model_path.exists():
                            model_path.unlink()
                            refresh_model_list()
                            messagebox.showinfo("成功", "模型已删除")
                    except Exception as e:
                        messagebox.showerror("错误", f"删除失败: {str(e)}")
        
        ttk.Button(button_frame, text="删除选中模型", 
                  command=delete_model).pack(side=tk.LEFT, padx=5)
        
        # 刷新按钮
        ttk.Button(button_frame, text="刷新列表", 
                  command=refresh_model_list).pack(side=tk.LEFT, padx=5)
        
        # 初始化显示
        refresh_model_list()
        if model_combo['values']:
            model_combo.set(model_combo['values'][0])

    def start_async_download(self, model_name, model_path):
        """启动异步下载的辅助方法"""
        def run_async():
            asyncio.run(self._download_model(model_name, model_path))
        
        threading.Thread(target=run_async).start()

    def toggle_custom_prompt(self):
        """切换自定义提示词输入框的状态"""
        if self.enable_custom_prompt.get():
            self.custom_prompt_text.config(state=tk.NORMAL)
        else:
            self.custom_prompt_text.delete('1.0', tk.END)
            self.custom_prompt_text.config(state=tk.DISABLED)

    def on_api_provider_change(self, event=None):
        """处理API提供商变更"""
        provider = self.api_provider_var.get()
        provider_config = self.api_providers[provider]
        
        if provider == "自定义":
            self.api_url_entry.config(state="normal")
            self.api_url_var.set("")
        else:
            self.api_url_entry.config(state="readonly")
            self.api_url_var.set(provider_config["base_url"])
        
        self.api_model_combo["values"] = provider_config["models"]
        if provider_config["models"]:
            self.api_model_var.set(provider_config["models"][0])
        else:
            self.api_model_var.set("")

    def refresh_models(self, model_type):
        """刷新模型列表"""
        if model_type == "local":
            models = self.scan_local_models()
            self.llm_model_combo["values"] = models
            if models:
                self.llm_model_var.set(models[0])
        else:
            provider = self.api_provider_var.get()
            if provider == "自定义":
                # 尝试从自定义API获取模型列表
                self.fetch_api_models()
            else:
                # 使用预定义的模型列表
                models = self.api_providers[provider]["models"]
                self.api_model_combo["values"] = models
                if models:
                    self.api_model_var.set(models[0])

    async def fetch_api_models(self):
        """从API获取可用模型列表"""
        try:
            url = self.api_url_var.get()
            api_key = self.api_key_var.get()
            
            if not url or not api_key:
                raise ValueError("请填写API URL和API Key")
            
            headers = {"Authorization": f"Bearer {api_key}"}
            
            async with aiohttp.ClientSession() as session:
                async with session.get(f"{url}/models", headers=headers) as response:
                    if response.status == 200:
                        data = await response.json()
                        # 根据不同API的响应格式处理
                        if "data" in data:  # OpenAI格式
                            models = [model["id"] for model in data["data"]]
                        elif "models" in data:  # 其他可能的格式
                            models = data["models"]
                        else:
                            models = []
                        
                        self.api_model_combo["values"] = models
                        if models:
                            self.api_model_var.set(models[0])
                    else:
                        raise Exception(f"API请求失败: {response.status}")
                        
        except Exception as e:
            logging.error(f"获取模型列表失败: {e}")
            messagebox.showerror("错误", f"获取模型列表失败: {str(e)}")

    async def optimize_with_llm(self, text, model_name):
        """使用LLM优化文本"""
        try:
            # 构建提示词
            prompt = self.build_prompt(text)
            
            if self.model_type_var.get() == "local":
                # 使用本地模型
                return await self.optimize_with_local_model(prompt, model_name)
            else:
                # 使用API模型
                return await self.optimize_with_api_model(prompt)
                
        except Exception as e:
            logging.error(f"LLM处理失败: {e}", exc_info=True)
            raise

    def build_prompt(self, text):
        """构建提示词"""
        if self.enable_custom_prompt.get():
            custom_prompt = self.custom_prompt_text.get('1.0', tk.END).strip()
            if custom_prompt:
                return custom_prompt + "\n\n" + text
        
        prompts = []
        if self.llm_optimize_options['improve_readability'].get():
            prompts.append("提高文本的可读性和流畅度")
        if self.llm_optimize_options['enhance_structure'].get():
            prompts.append("优化文本结构和段落组织")
        if self.llm_optimize_options['fix_grammar'].get():
            prompts.append("修正语法错误和表达不准确的地方")
        if self.llm_optimize_options['summarize'].get():
            prompts.append("为文本生成简短的摘要")
        if self.llm_optimize_options['translate'].get():
            prompts.append("将文本翻译成英文")
        
        prompt = "请" + "、".join(prompts) + "。以下是原文：\n\n" + text
        return prompt

    async def optimize_with_api_model(self, prompt):
        """使用API模型优化文本"""
        try:
            provider = self.api_provider_var.get()
            url = self.api_url_var.get()
            api_key = self.api_key_var.get()
            model = self.api_model_var.get()
            
            # 修改验证逻辑
            if not url:
                raise ValueError("请填写API URL")
            if not api_key:
                raise ValueError("请填写API Key")
            if not model and provider != "自定义":
                # 如果没有选择模型，使用默认的第一个模型
                available_models = self.api_providers[provider]["models"]
                if available_models:
                    model = available_models[0]
                    self.api_model_var.set(model)
                else:
                    raise ValueError("请选择模型")
            
            # 处理 API key
            if "," in api_key:
                # 如果包含多个 key，使用第一个
                api_key = api_key.split(",")[0].strip()
            
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            
            # 获取系统提示词
            if self.enable_custom_system_prompt.get():
                system_prompt = self.system_prompt_text.get('1.0', tk.END).strip()
                if not system_prompt:
                    system_prompt = self.default_system_prompt
            else:
                system_prompt = self.default_system_prompt

            # 构建API请求数据
            data = {
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                "stream": False,
                "max_tokens": self.api_params['max_tokens'].get(),
                "temperature": self.api_params['temperature'].get(),
                "top_p": self.api_params['top_p'].get(),
                "frequency_penalty": self.api_params['frequency_penalty'].get()
            }

            if provider == "Gitee AI":
                data.update({
                    "extra_body": {
                        "top_k": self.api_params['top_k'].get()
                    }
                })

            # 显示处理提示
            self.content_text.delete('1.0', tk.END)
            self.content_text.insert(tk.END, f"正在使用 {provider} 的 {model} 模型处理文本...\n")
            self.root.update()

            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{url}/chat/completions",
                    headers=headers,
                    json=data,
                    timeout=aiohttp.ClientTimeout(total=300)
                ) as response:
                    if response.status != 200:
                        error_data = await response.text()
                        raise Exception(f"API请求失败: {response.status}\n{error_data}")

                    result = await response.json()
                    if 'choices' in result and result['choices']:
                        content = result['choices'][0]['message']['content']
                        
                        # 清空处理提示
                        self.content_text.delete('1.0', tk.END)
                        
                        # 显示优化后的内容
                        self.content_text.insert(tk.END, "=== 优化结果 ===\n\n", "title")
                        self.content_text.insert(tk.END, content)
                        
                        # 配置标题样式
                        self.content_text.tag_configure("title", font=("Arial", 12, "bold"))
                        
                        return content
                    else:
                        raise Exception("API响应格式错误")

        except Exception as e:
            error_msg = f"API处理失败: {str(e)}"
            logging.error(error_msg, exc_info=True)
            
            # 显示错误信息
            self.content_text.delete('1.0', tk.END)
            self.content_text.insert(tk.END, error_msg + "\n", "error")
            self.content_text.tag_configure("error", foreground="red")
            
            # 显示错误对话框
            self.root.after(0, lambda: messagebox.showerror(
                "API处理失败",
                error_msg
            ))
            raise

    def toggle_system_prompt(self):
        """切换系统提示词编辑状态"""
        if self.enable_custom_system_prompt.get():
            self.system_prompt_text.configure(state=tk.NORMAL, fg='black')
        else:
            self.system_prompt_text.delete(1.0, tk.END)
            self.system_prompt_text.insert(tk.END, self.default_system_prompt)
            self.system_prompt_text.configure(state=tk.DISABLED, fg='gray')

    def reset_system_prompt(self):
        """重置系统提示词为默认值"""
        self.system_prompt_text.configure(state=tk.NORMAL)
        self.system_prompt_text.delete(1.0, tk.END)
        self.system_prompt_text.insert(tk.END, self.default_system_prompt)
        if not self.enable_custom_system_prompt.get():
            self.system_prompt_text.configure(state=tk.DISABLED, fg='gray')
        else:
            self.system_prompt_text.configure(fg='black')

    async def optimize_with_local_model(self, prompt, model_name):
        """使用本地模型优化文本"""
        try:
            from llama_cpp import Llama
            
            # 获取模型路径
            models_dir = Path("models")
            model_paths = {
                "chatglm3-6b": models_dir / "chatglm3-6b.Q4_K_M.gguf",
                "llama-2-7b": models_dir / "llama-2-7b.Q4_K_M.gguf",
                "qwen-7b": models_dir / "qwen-7b.Q4_K_M.gguf",
                "yi-6b": models_dir / "yi-6b.Q4_K_M.gguf",
                "mistral-7b": models_dir / "mistral-7b.Q4_K_M.gguf",
                "neural-7b": models_dir / "neural-7b.Q4_K_M.gguf"
            }
            
            # 规范化模型名称并获取模型路径
            normalized_model_name = model_name.lower().replace('_', '-').split('.')[0]
            model_path = next((path for key, path in model_paths.items() 
                              if key in normalized_model_name), None)
            
            if not model_path or not model_path.exists():
                raise FileNotFoundError(f"找不到模型文件: {model_path}")
            
            # 显示进度
            self.root.after(0, lambda: (
                self.progress_frame.pack(fill=tk.X, pady=5),
                self.progress_bar.pack(fill=tk.X),
                self.progress_label.pack(fill=tk.X)
            ))
            
            def update_progress(percentage, message):
                self.root.after(0, lambda: (
                    self.progress_var.set(percentage),
                    self.progress_label.config(text=message)
                ))
            
            update_progress(10, "加载模型中...")
            
            # 初始化模型
            llm = Llama(
                model_path=str(model_path),
                n_ctx=4096,  # 上下文窗口大小
                n_threads=max(1, os.cpu_count() // 2)  # 使用一半的CPU核心
            )
            
            update_progress(30, "模型加载完成，开始处理...")
            
            # 构建对话上下文
            messages = [
                {
                    "role": "system",
                    "content": "你是一个专业的文本优化助手，擅长提高文本的可读性、结构性和准确性。"
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ]
            
            # 生成回复
            update_progress(50, "正在生成优化内容...")
            
            response = llm.create_chat_completion(
                messages=messages,
                max_tokens=4096,
                temperature=0.7,
                top_p=0.9,
                stream=True
            )
            
            # 收集流式输出
            optimized_text = []
            current_progress = 50
            
            if isinstance(response, dict):
                # 非流式响应
                optimized_text = [response['choices'][0]['message']['content']]
            else:
                # 流式响应
                for chunk in response:
                    if 'choices' in chunk and chunk['choices']:
                        if 'delta' in chunk['choices'][0]:
                            content = chunk['choices'][0]['delta'].get('content', '')
                            if content:
                                optimized_text.append(content)
                                # 更新进度
                                current_progress = min(95, current_progress + 1)
                                update_progress(current_progress, "正在生成优化内容...")
            
            result = ''.join(optimized_text)
            
            update_progress(100, "优化完成!")
            self.root.after(3000, lambda: self.progress_frame.pack_forget())
            
            # 检查优化结果
            if not result or result.strip() == prompt.strip():
                raise ValueError("模型未能有效优化文本")
            
            return result
            
        except Exception as e:
            logging.error(f"本地模型处理失败: {e}", exc_info=True)
            self.root.after(0, lambda: messagebox.showerror(
                "优化失败",
                f"文本优化过程中发生错误：\n{str(e)}"
            ))
            raise

    def toggle_api_key_visibility(self):
        """切换 API Key 的显示/隐藏状态"""
        if self.show_api_key.get():
            self.api_key_entry.configure(show="")
        else:
            self.api_key_entry.configure(show="*")

    def switch_api_key(self):
        """切换使用的 API Key"""
        current_key = self.api_key_var.get()
        default_keys = [
            "99ZE2NVXCNLWIVWC6HQBGV5GMIKCEA9D8FXL16XN",
            "R6XZ3CRX2ZXWZ5XLCR3CLHDRNNQB6OAHYHTMJCU6"
        ]
        
        try:
            # 如果当前是组合的 keys，分割它们
            if "," in current_key:
                keys = current_key.split(",")
                # 使用第一个 key
                self.api_key_var.set(keys[0])
                if hasattr(self, 'key_indicator'):
                    self.key_indicator.config(text="使用: Key 1/2")
                return
            
            # 如果当前是单个 key，找到它在列表中的位置
            current_index = default_keys.index(current_key)
            # 切换到下一个 key（如果是最后一个，则回到第一个）
            next_index = (current_index + 1) % len(default_keys)
            self.api_key_var.set(default_keys[next_index])
            if hasattr(self, 'key_indicator'):
                self.key_indicator.config(text=f"使用: Key {next_index + 1}/2")
        except ValueError:
            # 如果当前的 key 不在默认列表中，使用第一个默认 key
            self.api_key_var.set(default_keys[0])
            if hasattr(self, 'key_indicator'):
                self.key_indicator.config(text="使用: Key 1/2")


if __name__ == "__main__":
    root = tk.Tk()
    app = CrawlerGUI(root)
    root.mainloop()